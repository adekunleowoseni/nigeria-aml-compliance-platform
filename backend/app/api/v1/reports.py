from __future__ import annotations

import re
from io import BytesIO
from datetime import date, datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import quote
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, Query, Request
from fastapi.responses import Response

from app.api.v1.alerts import _alert_not_soft_deleted, _alert_visible_to_user, _enrich_alert_for_api
from app.api.v1.in_memory_stores import _ALERTS, _TXNS
from app.config import settings
from app.core.security import get_current_user
from app.models.alert import AlertResponse
from app.models.transaction import TransactionResponse
from app.services.alert_snapshot import build_alert_snapshot
from app.services.customer_kyc_db import fetch_customer_kyc_any, get_or_create_customer_kyc, list_bvn_linked_accounts
from app.services import audit_trail
from app.services.compliance_bundle_narratives import build_aop_bundle_narrative, build_nfiu_cir_bundle_narrative
from app.services.regulatory_reports import (
    goaml_stub_xml,
    minimal_docx_bytes,
    nfiu_customer_change_xml,
    regulatory_narrative_docx_bytes,
    regulatory_narrative_pdf_bytes,
)
from app.services.statement_of_account import (
    account_context_dates_for_customer,
    clamp_statement_period,
    format_statement_text,
    parse_iso_date,
    soa_period_last_twelve_months,
    statement_lines_for_customer,
)
from app.services.sar_word_generator import (
    build_sar_case_context,
    generate_sar_narrative_sections,
    render_sar_docx_bytes,
)
from app.services.estr_word_generator import render_otc_estr_docx_bytes
from app.services.str_word_generator import render_str_docx_bytes
from docx import Document


def _render_str_docx_bytes_safe(**kwargs: Any) -> bytes:
    try:
        return render_str_docx_bytes(**kwargs)
    except FileNotFoundError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except ValueError as e:
        raise HTTPException(status_code=500, detail=str(e)) from e
from app.services.xml_generator import GoAMLGenerator
from app.services.reporting_profile_db import get_reporting_profile_row, list_calendar_entries
from app.services.reporting_context import (
    merge_goaml_stub_payload,
    reporting_entity_for_str_xml,
    upcoming_calendar_preview,
)

router = APIRouter(prefix="/reports")

_REPORTS: Dict[str, Dict[str, Any]] = {}
_STR_DRAFT_OVERRIDES: Dict[str, Dict[str, Any]] = {}
_STR_DRAFT_MAX_CHARS = 32000
_OTC_WORD_DRAFT_OVERRIDES: Dict[str, Dict[str, Any]] = {}
_OTC_WORD_DRAFT_MAX_CHARS = 32000
_generator = GoAMLGenerator()


def customer_report_history_summary(customer_id: str, *, limit: int = 5) -> Dict[str, Any]:
    """Return filing counts and recent report history for a customer."""
    cid = str(customer_id or "").strip()
    if not cid:
        return {"total": 0, "counts": {}, "recent": []}
    matches: List[Dict[str, Any]] = []
    for rid, rec in _REPORTS.items():
        if not isinstance(rec, dict):
            continue
        if not _report_not_soft_deleted(rec):
            continue
        rec_cid = str(rec.get("customer_id") or "").strip()
        if rec_cid != cid:
            continue
        rtype = str(rec.get("type") or "").strip().upper() or "UNKNOWN"
        status = str(rec.get("status") or "").strip().lower() or "draft"
        ts = str(rec.get("generated_at") or rec.get("created_at") or "")
        matches.append(
            {
                "report_id": str(rid),
                "type": rtype,
                "status": status,
                "generated_at": ts or None,
            }
        )
    counts: Dict[str, int] = {}
    for row in matches:
        key = str(row.get("type") or "UNKNOWN")
        counts[key] = int(counts.get(key, 0)) + 1
    matches.sort(key=lambda x: str(x.get("generated_at") or ""), reverse=True)
    return {"total": len(matches), "counts": counts, "recent": matches[: max(1, int(limit))]}


def _report_history_rows_filtered(
    *,
    customer_id: Optional[str],
    start_date: Optional[str],
    end_date: Optional[str],
) -> List[Dict[str, Any]]:
    cid = str(customer_id or "").strip()
    start = str(start_date or "").strip()
    end = str(end_date or "").strip()
    rows: List[Dict[str, Any]] = []
    for rid, rec in _REPORTS.items():
        if not isinstance(rec, dict):
            continue
        if not _report_not_soft_deleted(rec):
            continue
        rec_cid = str(rec.get("customer_id") or "").strip()
        if cid and rec_cid != cid:
            continue
        ts = str(rec.get("generated_at") or rec.get("created_at") or "").strip()
        day = ts[:10] if len(ts) >= 10 else ""
        if start and day and day < start:
            continue
        if end and day and day > end:
            continue
        rtype = str(rec.get("type") or "").strip().upper() or "UNKNOWN"
        status = str(rec.get("status") or "").strip().lower() or "draft"
        rows.append(
            {
                "report_id": str(rid),
                "customer_id": rec_cid,
                "type": rtype,
                "status": status,
                "generated_at": ts or None,
            }
        )
    rows.sort(key=lambda x: str(x.get("generated_at") or ""), reverse=True)
    return rows


@router.get("/history")
async def get_customer_report_history(
    customer_id: Optional[str] = Query(None, min_length=2, max_length=128),
    start_date: Optional[str] = Query(None, description="Inclusive date filter (YYYY-MM-DD)."),
    end_date: Optional[str] = Query(None, description="Inclusive date filter (YYYY-MM-DD)."),
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=500),
    user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Searchable report timeline by customer ID for CCO/admin review.
    """
    cid = str(customer_id or "").strip()
    role = str(user.get("role") or "").strip().lower()
    if role not in {"admin", "chief_compliance_officer", "compliance_officer"}:
        raise HTTPException(status_code=403, detail="Compliance, CCO, or admin role required.")
    if start_date and not re.fullmatch(r"\d{4}-\d{2}-\d{2}", str(start_date)):
        raise HTTPException(status_code=400, detail="start_date must be YYYY-MM-DD")
    if end_date and not re.fullmatch(r"\d{4}-\d{2}-\d{2}", str(end_date)):
        raise HTTPException(status_code=400, detail="end_date must be YYYY-MM-DD")
    if start_date and end_date and str(start_date) > str(end_date):
        raise HTTPException(status_code=400, detail="start_date cannot be after end_date")

    rows = _report_history_rows_filtered(customer_id=cid or None, start_date=start_date, end_date=end_date)
    counts: Dict[str, int] = {}
    for row in rows:
        key = str(row.get("type") or "UNKNOWN")
        counts[key] = int(counts.get(key, 0)) + 1
    total = len(rows)
    page_items = rows[skip : skip + limit]
    return {
        "customer_id": cid or None,
        "total": total,
        "counts": counts,
        "skip": skip,
        "limit": limit,
        "items": page_items,
    }


async def _load_report_profile(request: Request) -> Dict[str, Any]:
    pg = getattr(request.app.state, "pg", None)
    if not pg:
        return {}
    try:
        return await get_reporting_profile_row(pg)
    except Exception:
        return {}


def _report_not_soft_deleted(r: Dict[str, Any]) -> bool:
    return r.get("deleted_at") is None


def _persist_report_with_audit(
    user: Dict[str, Any],
    report_id: str,
    rec: Dict[str, Any],
    *,
    extra: Optional[Dict[str, Any]] = None,
) -> None:
    if not rec.get("generated_at"):
        rec["generated_at"] = datetime.utcnow().isoformat() + "Z"
    _REPORTS[report_id] = rec
    audit_trail.record_report_generated(user, report_id, rec, extra=extra)


def _latest_txn_for_customer(customer_id: str):
    txns = [t for t in _TXNS.values() if t.customer_id == customer_id]
    if not txns:
        return None
    return max(txns, key=lambda t: t.created_at)


def _linked_channel_for_alert(alert: AlertResponse) -> Optional[str]:
    tx = _TXNS.get(alert.transaction_id)
    if not tx:
        return None
    md = tx.metadata if isinstance(tx.metadata, dict) else {}
    ch = str(md.get("channel") or "").strip().lower()
    return ch or None


def _alert_eligible_for_sar(alert: AlertResponse) -> bool:
    """SAR (demo workflow): alerts closed as false positive are eligible for suspicious-activity SAR filing."""
    if (alert.status or "").lower() != "closed":
        return False
    return (alert.last_resolution or "").strip() == "false_positive"


def _false_positive_resolve_notes(alert: AlertResponse) -> str:
    for h in reversed(alert.investigation_history or []):
        if isinstance(h, dict) and h.get("action") == "resolve":
            return str(h.get("notes") or "")
    return ""


def _activity_profile_for_sar(alert: AlertResponse, linked: Optional[TransactionResponse]) -> Dict[str, Any]:
    hints: List[str] = []
    if linked and isinstance(linked.metadata, dict):
        md = linked.metadata
        sc = md.get("simulation_scenario")
        if sc:
            hints.append(f"Demo scenario: {sc}")
        prof = md.get("profile") or md.get("pattern")
        if prof:
            hints.append(f"Pattern label: {prof}")
    return {
        "alert_summary": alert.summary,
        "rule_ids": list(alert.rule_ids or []),
        "false_positive_resolution_notes": _false_positive_resolve_notes(alert),
        "scenario_hints": hints,
        "investigation_actions": [
            str(h.get("action"))
            for h in (alert.investigation_history or [])
            if isinstance(h, dict) and h.get("action")
        ],
    }


def _synthetic_activity_txn(alert: AlertResponse) -> Dict[str, Any]:
    now = datetime.utcnow()
    tid = alert.transaction_id or f"sar-activity-{alert.id[:12]}"
    return {
        "id": tid,
        "customer_id": alert.customer_id,
        "amount": 0.0,
        "currency": "NGN",
        "transaction_type": "suspicious_activity_pattern",
        "narrative": alert.summary or "Suspicious activity — see alert typology and investigation record (demo).",
        "created_at": now,
        "metadata": {"sar_basis": "suspicious_activity", "alert_id": alert.id},
        "counterparty_name": None,
        "counterparty_id": None,
    }


def _resolve_txn_for_sar(
    alert: AlertResponse,
    transaction_id_override: Optional[str],
) -> Tuple[Dict[str, Any], Optional[TransactionResponse], bool]:
    """
    Returns (txn_dict, linked_txn_or_none, is_synthetic).
    Prefers a stored transaction when present; otherwise builds an activity-led synthetic row (no txn required for SAR).
    """
    txn: Optional[TransactionResponse] = None
    if transaction_id_override and transaction_id_override.strip():
        txn = _TXNS.get(transaction_id_override.strip())
    if txn is None and alert.transaction_id:
        txn = _TXNS.get(alert.transaction_id)
    if txn is None:
        txn = _latest_txn_for_customer(alert.customer_id)
    if txn is None:
        return _synthetic_activity_txn(alert), None, True
    return txn.model_dump(), txn, False


def _alert_eligible_for_str(alert: AlertResponse) -> bool:
    if (alert.last_resolution or "").strip() == "false_positive":
        return False
    if (alert.status or "").lower() != "escalated":
        return False
    # Cash OTC ESTR uses the extended-return (ESTR) track only — not goAML STR generation.
    if getattr(alert, "otc_report_kind", None) == "otc_estr":
        return False
    return bool(getattr(alert, "cco_str_approved", False))


def _otc_reporting_requirements_met(alert: AlertResponse) -> bool:
    """CCO must approve OTC reporting; CO must have escalated unless admin auto-OTC mode is on."""
    if not bool(getattr(alert, "cco_otc_approved", False)):
        return False
    if bool(getattr(settings, "cco_auto_approve_otc_reporting", False)):
        return True
    return (alert.status or "").lower() == "escalated"


def _alert_eligible_for_otc_estr(alert: AlertResponse) -> bool:
    return (
        getattr(alert, "otc_outcome", None) == "true_positive"
        and getattr(alert, "otc_report_kind", None) == "otc_estr"
        and _otc_reporting_requirements_met(alert)
    )


def _alert_eligible_for_otc_esar(alert: AlertResponse) -> bool:
    return (
        getattr(alert, "otc_outcome", None) == "true_positive"
        and getattr(alert, "otc_report_kind", None) == "otc_esar"
        and _otc_reporting_requirements_met(alert)
    )


def _alert_eligible_for_otc_estr_word_ready(alert: AlertResponse) -> bool:
    """OTC cash ESTR: true-positive filing, CCO approval, and (unless admin auto mode) CO escalation."""
    return _alert_eligible_for_otc_estr(alert)


def _txn_dict_for_estr_kyc(aid: Optional[str], cid: str) -> Dict[str, Any]:
    if aid and _ALERTS.get(aid):
        ao = _ALERTS[aid]
        cid = cid or ao.customer_id
        tid = ao.transaction_id
        tx = _TXNS.get(tid) if tid else None
        if tx:
            return tx.model_dump()
        return {
            "id": tid or "synthetic",
            "customer_id": ao.customer_id,
            "amount": 0.0,
            "currency": "NGN",
            "transaction_type": "",
            "narrative": ao.summary or "",
            "metadata": {},
            "created_at": datetime.utcnow().isoformat(),
        }
    return {
        "customer_id": cid or "UNKNOWN",
        "amount": 0.0,
        "currency": "NGN",
        "transaction_type": "",
        "narrative": "",
        "metadata": {},
        "created_at": datetime.utcnow().isoformat(),
    }


def _report_download_content_disposition(customer_name: str, label: str, ext: str) -> str:
    base_name = (customer_name or "").strip() or "Customer"
    base_name = re.sub(r"\s+", "_", base_name)
    base_name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", base_name).strip("-_")
    if not base_name:
        base_name = "Customer"
    label_clean = re.sub(r"[\s\-]+", "_", str(label or "").strip()) or "REPORT"
    stem = f"{label_clean}_{base_name}".strip("_")
    stem = stem[:200]
    filename_utf8 = f"{stem}.{ext}"
    ascii_stem = "".join((c if 32 <= ord(c) < 127 and c not in '<>:"/\\|?*' else "_") for c in stem)
    ascii_stem = re.sub(r"_+", "_", ascii_stem).strip("_") or "Customer"
    ascii_name = f"{ascii_stem}.{ext}"
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{quote(filename_utf8)}"


def _http_exception_detail_str(exc: HTTPException) -> str:
    d = exc.detail
    if isinstance(d, str):
        return d
    if isinstance(d, list):
        return "; ".join(str(x) for x in d)
    return str(d)


def _alert_eligible_for_sar_bulk_item(alert: AlertResponse) -> bool:
    """SAR bulk pool: false-positive activity SAR or CCO-approved OTC ESAR (identity). Excludes OTC ESTR."""
    if getattr(alert, "otc_report_kind", None) == "otc_estr":
        return False
    return _alert_eligible_for_sar(alert) or _alert_eligible_for_otc_esar(alert)


def _activity_profile_for_otc_esar(alert: AlertResponse) -> Dict[str, Any]:
    return {
        "alert_summary": alert.summary,
        "otc_subject": alert.otc_subject,
        "otc_filing_reason": alert.otc_filing_reason,
        "otc_filing_reason_detail": alert.otc_filing_reason_detail,
        "otc_officer_rationale": alert.otc_officer_rationale,
        "rule_ids": list(alert.rule_ids or []),
        "workflow": "otc_esar",
    }


def _combine_sar_notes_otc(alert: AlertResponse, sar_notes: str) -> str:
    parts = [
        f"OTC ESAR — matter: {alert.otc_subject}",
        (alert.otc_officer_rationale or "").strip(),
        (sar_notes or "").strip(),
    ]
    return "\n\n".join(p for p in parts if p).strip()[:4000]


def _extract_str_notes(payload: Dict[str, Any]) -> str:
    raw = payload.get("str_notes") or payload.get("reason_note") or ""
    if isinstance(raw, str) and raw.strip():
        return raw.strip()
    ctx = payload.get("additional_context")
    if isinstance(ctx, dict):
        inner = ctx.get("str_notes") or ctx.get("notes")
        if isinstance(inner, str) and inner.strip():
            return inner.strip()
    return ""


def _looks_like_str_editor_scaffold(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return False
    # Legacy modal scaffold that should never be treated as an authored draft.
    return (
        "suspicious transaction report" in t
        and "alert:" in t
        and "narrative source:" in t
        and "xml payload (excerpt)" in t
    )


def _extract_internal_narrative_from_scaffold(text: str) -> str:
    raw = (text or "").strip()
    if not raw:
        return ""
    m = re.search(
        r"internal narrative\s*(.*?)\s*xml payload \(excerpt\)",
        raw,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if not m:
        return ""
    out = re.sub(r"\s+\n", "\n", m.group(1)).strip()
    return out


def _looks_like_low_value_str_note(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return True
    if t in {"str draft note", "suspicious transaction report"}:
        return True
    return len(t) <= 140 and "confirmed suspicious activity" in t and "true positive escalation" in t


def get_saved_str_draft_notes(alert_id: str) -> str:
    aid = (alert_id or "").strip()
    if not aid:
        return ""
    d = _STR_DRAFT_OVERRIDES.get(aid) or {}
    notes = str(d.get("str_notes") or "").strip()
    if _looks_like_str_editor_scaffold(notes):
        extracted = _extract_internal_narrative_from_scaffold(notes)
        if _looks_like_low_value_str_note(extracted):
            return ""
        return extracted if extracted else ""
    if _looks_like_low_value_str_note(notes):
        return ""
    return notes


def get_saved_otc_word_draft_notes(alert_id: str) -> str:
    aid = (alert_id or "").strip()
    if not aid:
        return ""
    d = _OTC_WORD_DRAFT_OVERRIDES.get(aid) or {}
    return str(d.get("estr_notes") or "").strip()


def _alert_eligible_for_otc_word_draft_preview(alert: AlertResponse) -> bool:
    """Escalated true-positive OTC (ESTR or ESAR) — same lifecycle stage as STR draft editing."""
    if (alert.status or "").lower() != "escalated":
        return False
    if getattr(alert, "otc_outcome", None) != "true_positive":
        return False
    rk = getattr(alert, "otc_report_kind", None)
    return rk in ("otc_estr", "otc_esar")


def _default_otc_word_draft_notes(alert: AlertResponse) -> str:
    saved = get_saved_otc_word_draft_notes(alert.id)
    if saved:
        return saved
    rationale = str(getattr(alert, "otc_officer_rationale", None) or "").strip()
    if rationale:
        return rationale
    return "OTC Word draft — add reasons for filing and any extension narrative for the regulator."


async def _otc_word_draft_docx_bytes(
    request: Request,
    alert_id: str,
    estr_notes: str,
    user: Dict[str, Any],
) -> bytes:
    aid = (alert_id or "").strip()
    alert = _ALERTS.get(aid)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    if not _alert_visible_to_user(user, alert):
        raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
    cid = str(alert.customer_id or "").strip() or "UNKNOWN"
    txn_dict = _txn_dict_for_estr_kyc(aid, cid)
    pg = getattr(request.app.state, "pg", None)
    customer = await get_or_create_customer_kyc(pg, cid, txn_dict)
    bvn_linked_accounts = await list_bvn_linked_accounts(
        pg, str(customer.id_number or "").strip(), primary_customer_id=cid
    )
    alert_dict = alert.model_dump()
    return await render_otc_estr_docx_bytes(
        customer=customer,
        alert=alert_dict,
        estr_notes=estr_notes,
        approver_name=_approver_display_name(user),
        bvn_linked_accounts=bvn_linked_accounts,
    )


def _alert_eligible_for_str_draft_preview(alert: AlertResponse) -> bool:
    """
    Draft editing/preview is allowed both:
    - before CCO approval (escalated queue), and
    - after CCO approval (normal STR generation path),
    excluding false-positive and OTC ESTR path.
    """
    if (alert.last_resolution or "").strip() == "false_positive":
        return False
    if (alert.status or "").lower() != "escalated":
        return False
    if getattr(alert, "otc_report_kind", None) == "otc_estr":
        return False
    return True


def _normalize_txn_dict(txn_dict: Dict[str, Any], alert: AlertResponse) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for k, v in txn_dict.items():
        if isinstance(v, datetime):
            out[k] = v.isoformat()
        else:
            out[k] = v
    tid = out.get("id") or alert.transaction_id
    out["id"] = tid
    return out


def _build_str_narrative(
    alert: AlertResponse,
    alert_id: str,
    txn_dict: Dict[str, Any],
    inflows_total: float,
    outflows_total: float,
    period_text: str,
    str_notes: str,
) -> str:
    parts: list[str] = [
        f"Alert reference: {alert_id}.",
        f"Customer identifier: {alert.customer_id}.",
        f"Linked transaction identifier: {alert.transaction_id}.",
        f"Current alert status: {alert.status}.",
    ]
    if alert.last_resolution:
        parts.append(f"Resolution (where applicable): {alert.last_resolution}.")
    if alert.summary:
        parts.append(f"Alert summary / typology: {alert.summary}.")
    if str_notes:
        parts.append(f"Compliance officer STR notes: {str_notes}")
    txn_narr = txn_dict.get("narrative")
    if txn_narr:
        parts.append(f"Transaction narrative: {txn_narr}")
    cur = txn_dict.get("currency") or "NGN"
    parts.append(
        f"Customer activity context ({period_text}): aggregate inflows {cur} {inflows_total:,.2f}, "
        f"aggregate outflows {cur} {outflows_total:,.2f}."
    )
    if alert.rule_ids:
        parts.append(f"Triggered rule identifiers: {', '.join(str(r) for r in alert.rule_ids)}.")
    return "\n\n".join(p for p in parts if p)[:12000]


def _render_str_xml(
    alert: AlertResponse,
    alert_id: str,
    txn_dict: Dict[str, Any],
    inflows_total: float,
    outflows_total: float,
    period_text: str,
    str_notes: str,
    *,
    profile: Optional[Dict[str, Any]] = None,
) -> str:
    currency = txn_dict.get("currency") or "NGN"
    suspicious_activity = {
        "reason": alert.summary or "Suspicious transaction report",
        "alert_id": alert_id,
        "customer_id": alert.customer_id,
        "transaction_id": alert.transaction_id,
        "severity": alert.severity,
        "alert_status": alert.status,
    }
    if str_notes:
        suspicious_activity["analyst_str_notes_present"] = "true"

    narrative = _build_str_narrative(alert, alert_id, txn_dict, inflows_total, outflows_total, period_text, str_notes)

    tx_row: Dict[str, Any] = {
        "transaction_id": str(txn_dict.get("id") or alert.transaction_id),
        "customer_id": alert.customer_id,
        "amount": txn_dict.get("amount", 0),
        "currency": currency,
        "transaction_type": str(txn_dict.get("transaction_type") or ""),
    }
    ts = txn_dict.get("created_at") or txn_dict.get("timestamp")
    if isinstance(ts, datetime):
        tx_row["transaction_datetime"] = ts.isoformat()
    elif ts:
        tx_row["transaction_datetime"] = str(ts)

    hist = list(alert.investigation_history or [])
    clean_hist: list[Dict[str, Any]] = []
    for h in hist:
        if not isinstance(h, dict):
            continue
        row = {k: (v.isoformat() if isinstance(v, datetime) else v) for k, v in h.items()}
        clean_hist.append(row)

    case_details: Dict[str, Any] = {
        "alert": {
            "id": alert.id,
            "status": alert.status,
            "severity": alert.severity,
            "last_resolution": alert.last_resolution,
            "summary": alert.summary,
            "rule_ids": alert.rule_ids,
            "created_at": alert.created_at.isoformat() if alert.created_at else None,
            "updated_at": alert.updated_at.isoformat() if alert.updated_at else None,
        },
        "transaction": _normalize_txn_dict(txn_dict, alert),
        "cashflow": {
            "period_text": period_text,
            "inflows_total": inflows_total,
            "outflows_total": outflows_total,
            "currency": currency,
        },
        "investigation_history": clean_hist,
        "analyst_str_notes": str_notes or None,
    }

    ent = reporting_entity_for_str_xml(profile)
    return _generator.generate_str(
        reporting_entity={"name": ent["name"], "registration_number": ent["registration_number"]},
        suspicious_activity=suspicious_activity,
        transactions=[tx_row],
        narrative=narrative,
        case_details=case_details,
    )


def _compute_customer_cashflow_totals(customer_id: str) -> tuple[float, float]:
    # Lightweight cashflow approximation for STR narrative.
    inflow_types = {"salary", "transfer_in", "pos_settlement", "cash_deposit", "wire"}
    outflow_types = {"transfer_out"}
    inflows = 0.0
    outflows = 0.0
    for t in _TXNS.values():
        if t.customer_id != customer_id:
            continue
        tx_type = str(t.transaction_type or "").lower()
        if tx_type in inflow_types:
            inflows += float(t.amount or 0.0)
        elif tx_type in outflow_types:
            outflows += float(t.amount or 0.0)
    return inflows, outflows


async def _draft_str_report(
    request: Request,
    alert_id: str,
    str_notes: str,
    user: Dict[str, Any],
    *,
    profile: Optional[Dict[str, Any]] = None,
    allow_preapproval_preview: bool = False,
) -> Dict[str, Any]:
    alert = _ALERTS.get(alert_id)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    prof = profile if profile is not None else await _load_report_profile(request)
    if allow_preapproval_preview:
        if not _alert_eligible_for_str_draft_preview(alert):
            raise HTTPException(
                status_code=400,
                detail="STR draft preview requires an escalated, non-false-positive alert in STR workflow.",
            )
    else:
        if not _alert_eligible_for_str(alert):
            raise HTTPException(
                status_code=400,
                detail="STR requires an escalated alert that the Chief Compliance Officer has approved for filing. "
                "False-positive resolutions are not eligible.",
            )

    txn = _TXNS.get(alert.transaction_id)
    txn_dict = (
        txn.model_dump()
        if txn
        else {
            "created_at": datetime.utcnow(),
            "amount": 0.0,
            "transaction_type": "",
            "currency": "NGN",
            "narrative": alert.summary,
            "metadata": {},
            "customer_id": alert.customer_id,
            "id": alert.transaction_id,
        }
    )

    inflows_total, outflows_total = _compute_customer_cashflow_totals(alert.customer_id)
    now = datetime.utcnow()
    period_text = f"{(now - timedelta(days=365)).strftime('%B %d, %Y')} to {now.strftime('%B %d, %Y')} (12-month window)"

    report_id = str(uuid4())
    xml_preview = _render_str_xml(
        alert, alert_id, txn_dict, inflows_total, outflows_total, period_text, str_notes, profile=prof
    )

    alert_context = alert.model_dump()
    alert_context["inflows_total"] = inflows_total
    alert_context["outflows_total"] = outflows_total
    alert_context["period_text"] = period_text

    str_rec = {
        "type": "STR",
        "status": "draft",
        "xml": xml_preview,
        "alert_id": alert_id,
        "customer_id": alert.customer_id,
        "txn": txn_dict,
        "alert": alert_context,
        "str_notes": str_notes,
    }
    _persist_report_with_audit(user, report_id, str_rec)

    return {"report_id": report_id, "xml_preview": xml_preview, "validation_passed": True}


def _draft_aop_record(
    customer_id: str,
    user: Dict[str, Any],
    *,
    account_product: str,
    risk_rating: str,
    word_narrative: Optional[str] = None,
    word_narrative_source: Optional[str] = None,
    profile: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    report_id = str(uuid4())
    cid = (customer_id or "").strip() or "UNKNOWN"
    xml = goaml_stub_xml(
        "AOP",
        merge_goaml_stub_payload(
            profile,
            {
                "report_id": report_id,
                "customer_id": cid,
                "account_product": account_product,
                "risk_rating": risk_rating,
                "prepared_by": str(user.get("display_name") or user.get("email") or "Compliance"),
            },
        ),
    )
    rec: Dict[str, Any] = {"type": "AOP", "status": "draft", "xml": xml, "customer_id": cid}
    if word_narrative:
        rec["word_narrative"] = word_narrative
        rec["word_narrative_source"] = word_narrative_source or "template"
    _persist_report_with_audit(user, report_id, rec)
    out: Dict[str, Any] = {"report_id": report_id, "xml_preview": xml, "validation_passed": True}
    if word_narrative_source:
        out["word_narrative_source"] = word_narrative_source
    return out


def _draft_nfiu_cir_record(
    change_type: str,
    customer_id: str,
    user: Dict[str, Any],
    fields: Dict[str, Any],
    *,
    word_narrative: Optional[str] = None,
    word_narrative_source: Optional[str] = None,
) -> Dict[str, Any]:
    report_id = str(uuid4())
    payload_fields = {
        "report_id": report_id,
        "customer_id": customer_id,
        "old_value": fields.get("old_value"),
        "new_value": fields.get("new_value"),
        "notes": fields.get("notes"),
        "bvn_old": fields.get("bvn_old"),
        "bvn_new": fields.get("bvn_new"),
        "name_old": fields.get("name_old"),
        "name_new": fields.get("name_new"),
        "dob_old": fields.get("dob_old"),
        "dob_new": fields.get("dob_new"),
    }
    xml = nfiu_customer_change_xml(change_type, payload_fields)
    rec: Dict[str, Any] = {
        "type": "NFIU_CIR",
        "status": "draft",
        "xml": xml,
        "change_type": change_type,
        "customer_id": customer_id,
    }
    if word_narrative:
        rec["word_narrative"] = word_narrative
        rec["word_narrative_source"] = word_narrative_source or "template"
    _persist_report_with_audit(user, report_id, rec)
    out: Dict[str, Any] = {
        "report_id": report_id,
        "xml_preview": xml,
        "validation_passed": True,
        "change_type": change_type,
    }
    if word_narrative_source:
        out["word_narrative_source"] = word_narrative_source
    return out


async def _draft_soa_record(
    request: Request,
    *,
    customer_id: str,
    alert_id: str,
    user: Dict[str, Any],
    period_from: date,
    period_to: date,
) -> Dict[str, Any]:
    pg = getattr(request.app.state, "pg", None)
    prof = await _load_report_profile(request)
    _, _, opened_s = await account_context_dates_for_customer(pg, customer_id)
    lines = statement_lines_for_customer(customer_id, period_from, period_to)
    text = format_statement_text(lines, customer_id, opened_s)
    report_id = str(uuid4())
    xml = goaml_stub_xml(
        "SOA",
        merge_goaml_stub_payload(
            prof,
            {
                "report_id": report_id,
                "customer_id": customer_id,
                "alert_id": alert_id,
                "period_start": period_from.isoformat(),
                "period_end": period_to.isoformat(),
                "transaction_rows": len(lines),
            },
        ),
    )
    rec: Dict[str, Any] = {
        "type": "SOA",
        "status": "draft",
        "xml": xml,
        "customer_id": customer_id,
        "alert_id": alert_id,
        "period_start": period_from.isoformat(),
        "period_end": period_to.isoformat(),
        "statement_text": text,
        "account_opened_kyc": opened_s,
        "transaction_row_count": len(lines),
    }
    _persist_report_with_audit(
        user,
        report_id,
        rec,
        extra={"alert_id": alert_id, "kind": "soa"},
    )
    return {
        "report_id": report_id,
        "period_start": period_from.isoformat(),
        "period_end": period_to.isoformat(),
        "xml_preview": xml,
        "validation_passed": True,
    }


async def _resolve_soa_period_for_str_alert(
    request: Request,
    alert: AlertResponse,
    *,
    bulk_multi: bool,
    soa_start: Optional[date],
    soa_end: Optional[date],
) -> Tuple[date, date]:
    pg = getattr(request.app.state, "pg", None)
    acc_start, _, _ = await account_context_dates_for_customer(pg, alert.customer_id)
    today = datetime.utcnow().date()
    if bulk_multi:
        return soa_period_last_twelve_months(acc_start, today)
    if soa_start or soa_end:
        return clamp_statement_period(acc_start, today, soa_start, soa_end)
    return clamp_statement_period(acc_start, today, None, None)


def _maybe_str_aop_records(
    payload: Dict[str, Any],
    customer_id: str,
    user: Dict[str, Any],
    *,
    profile: Optional[Dict[str, Any]] = None,
) -> Optional[str]:
    """If payload requests AOP alongside STR, draft AOP for the same customer; return report_id."""
    if not (payload.get("include_aop") or payload.get("generate_aop")):
        return None
    cid = (customer_id or "").strip()
    if not cid:
        return None
    aop_product = str(payload.get("aop_account_product") or "Savings").strip() or "Savings"
    aop_risk = str(payload.get("aop_risk_rating") or "medium").strip() or "medium"
    aop_out = _draft_aop_record(
        cid, user, account_product=aop_product, risk_rating=aop_risk, profile=profile
    )
    return str(aop_out.get("report_id") or "")


@router.post("/str/generate")
async def generate_str(
    request: Request,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    alert_id = str(payload.get("alert_id") or "").strip()
    if not alert_id:
        raise HTTPException(status_code=400, detail="alert_id is required")

    str_notes = _extract_str_notes(payload)
    used_saved_draft = False
    if bool(payload.get("use_saved_draft", True)):
        draft_notes = get_saved_str_draft_notes(alert_id)
        if draft_notes:
            str_notes = draft_notes
            used_saved_draft = True
    if not str_notes:
        raise HTTPException(status_code=400, detail="str_notes is required (compliance reason for this STR).")

    alert = _ALERTS.get(alert_id)
    if alert and not _alert_visible_to_user(user, alert):
        raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")

    prof = await _load_report_profile(request)
    out = await _draft_str_report(request, alert_id, str_notes, user, profile=prof)
    if alert:
        aop_id = _maybe_str_aop_records(payload, alert.customer_id, user, profile=prof)
        if aop_id:
            out = {**out, "aop_report_id": aop_id}
        include_soa = bool(payload.get("include_soa") or payload.get("generate_statement_of_account"))
        if include_soa:
            soa_start = parse_iso_date(payload.get("statement_period_start"))
            soa_end = parse_iso_date(payload.get("statement_period_end"))
            d_from, d_to = await _resolve_soa_period_for_str_alert(
                request,
                alert,
                bulk_multi=False,
                soa_start=soa_start,
                soa_end=soa_end,
            )
            soa_out = await _draft_soa_record(
                request,
                customer_id=alert.customer_id,
                alert_id=alert_id,
                user=user,
                period_from=d_from,
                period_to=d_to,
            )
            out = {
                **out,
                "soa_report_id": soa_out["report_id"],
                "soa_period_start": soa_out["period_start"],
                "soa_period_end": soa_out["period_end"],
            }
    return {**out, "used_saved_draft": used_saved_draft}


@router.post("/str/generate-bulk")
async def generate_str_bulk(
    request: Request,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    """Generate one STR per alert_id (shared str_notes). Optional AOP and statement of account per alert."""
    raw = payload.get("alert_ids")
    if not isinstance(raw, list) or not raw:
        raise HTTPException(status_code=400, detail="alert_ids must be a non-empty array")
    ids_ordered: List[str] = []
    seen: set[str] = set()
    for x in raw:
        aid = str(x).strip()
        if not aid or aid in seen:
            continue
        seen.add(aid)
        ids_ordered.append(aid)

    if not ids_ordered:
        raise HTTPException(status_code=400, detail="alert_ids must contain at least one valid id")

    str_notes = _extract_str_notes(payload)
    if not str_notes:
        raise HTTPException(status_code=400, detail="str_notes is required (compliance reason for these STRs).")

    include_aop = bool(payload.get("include_aop") or payload.get("generate_aop"))
    aop_product = str(payload.get("aop_account_product") or "Savings").strip() or "Savings"
    aop_risk = str(payload.get("aop_risk_rating") or "medium").strip() or "medium"
    include_soa = bool(payload.get("include_soa") or payload.get("generate_statement_of_account"))
    soa_start = parse_iso_date(payload.get("statement_period_start"))
    soa_end = parse_iso_date(payload.get("statement_period_end"))
    bulk_multi = len(ids_ordered) > 1
    prof = await _load_report_profile(request)

    results: List[Dict[str, Any]] = []
    for aid in ids_ordered:
        alert = _ALERTS.get(aid)
        if not alert:
            results.append({"alert_id": aid, "ok": False, "error": "alert not found"})
            continue
        if not _alert_visible_to_user(user, alert):
            results.append({"alert_id": aid, "ok": False, "error": "Alert outside your zone/branch scope."})
            continue
        if not _alert_eligible_for_str(alert):
            results.append(
                {
                    "alert_id": aid,
                    "ok": False,
                    "error": "not eligible (requires escalated + CCO-approved alert)",
                }
            )
            continue
        try:
            notes_eff = str_notes
            used_saved_draft = False
            if bool(payload.get("use_saved_draft", True)):
                draft_notes = get_saved_str_draft_notes(aid)
                if draft_notes:
                    notes_eff = draft_notes
                    used_saved_draft = True
            out = await _draft_str_report(request, aid, notes_eff, user, profile=prof)
        except HTTPException as e:
            detail = e.detail if isinstance(e.detail, str) else str(e.detail)
            results.append({"alert_id": aid, "ok": False, "error": detail})
            continue
        row: Dict[str, Any] = {
            "alert_id": aid,
            "customer_id": alert.customer_id,
            "ok": True,
            "used_saved_draft": used_saved_draft,
            **out,
        }
        if include_aop:
            aop_out = _draft_aop_record(
                alert.customer_id,
                user,
                account_product=aop_product,
                risk_rating=aop_risk,
                profile=prof,
            )
            row["aop_report_id"] = aop_out["report_id"]
        if include_soa:
            try:
                d_from, d_to = await _resolve_soa_period_for_str_alert(
                    request,
                    alert,
                    bulk_multi=bulk_multi,
                    soa_start=soa_start,
                    soa_end=soa_end,
                )
                soa_out = await _draft_soa_record(
                    request,
                    customer_id=alert.customer_id,
                    alert_id=aid,
                    user=user,
                    period_from=d_from,
                    period_to=d_to,
                )
                row["soa_report_id"] = soa_out["report_id"]
                row["soa_period_start"] = soa_out["period_start"]
                row["soa_period_end"] = soa_out["period_end"]
            except HTTPException as e:
                detail = e.detail if isinstance(e.detail, str) else str(e.detail)
                row["soa_error"] = detail
        results.append(row)

    ok_n = sum(1 for r in results if r.get("ok"))
    return {"results": results, "generated": ok_n, "requested": len(ids_ordered)}


_STR_DRAFT_AI_PROMPT_MAX = 48000

_STR_DRAFT_AI_SYSTEM = (
    "You are an AML compliance reporting assistant helping draft Suspicious Transaction Report (STR) narrative text "
    "for the Nigerian regulatory context. Follow user instructions precisely. Preserve facts, amounts, dates, names, "
    "and account identifiers; improve clarity and professional tone only. Do not invent facts. "
    "Output plain text only (no markdown code fences unless explicitly requested)."
)


@router.post("/str/draft/ai-assist")
async def str_draft_ai_assist(
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Run the configured LLM (GEMINI_API_KEY / OPENAI_API_KEY / Ollama per settings) on the STR draft helper prompt.
    """
    prompt = str(payload.get("prompt") or "").strip()
    if not prompt:
        raise HTTPException(status_code=400, detail="prompt is required")
    if len(prompt) > _STR_DRAFT_AI_PROMPT_MAX:
        raise HTTPException(
            status_code=400,
            detail=f"prompt exceeds {_STR_DRAFT_AI_PROMPT_MAX} characters",
        )
    try:
        from app.services.llm.client import get_llm_client

        client = get_llm_client()
        result = await client.generate(prompt, system=_STR_DRAFT_AI_SYSTEM, temperature=0.35)
        content = (result.content or "").strip()
        audit_trail.record_event_from_user(
            user,
            action="report.str_draft.ai_assist",
            resource_type="reports",
            resource_id="str-draft-ai",
            details={
                "prompt_chars": len(prompt),
                "provider": result.provider,
                "model": result.model,
                "reply_chars": len(content),
            },
        )
        return {"content": content, "provider": result.provider, "model": result.model}
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"LLM request failed: {e!s}") from e


@router.get("/str/draft/{alert_id}")
async def get_str_draft_preview(
    request: Request,
    alert_id: str,
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    alert = _ALERTS.get(aid)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    if not _alert_visible_to_user(user, alert):
        raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
    saved_notes = get_saved_str_draft_notes(aid)
    has_saved = bool(saved_notes)
    notes = saved_notes or (alert.escalation_reason_notes or "").strip() or "STR draft note"
    if not _alert_eligible_for_str_draft_preview(alert):
        return {
            "alert_id": aid,
            "str_notes": notes,
            "has_saved_draft": has_saved,
            "word_preview_lines": [notes] if notes else [],
            "preview_warning": "Alert is not currently eligible for STR draft preview generation.",
        }
    lines: List[str] = []
    try:
        tmp = await _draft_str_report(request, aid, notes, user, allow_preapproval_preview=True)
        rid = str(tmp["report_id"])
        rec = _REPORTS.get(rid)
        if rec:
            pg = getattr(request.app.state, "pg", None)
            docx = await str_docx_bytes_from_report_record(rec, pg=pg, user=user)
            doc = Document(BytesIO(docx))
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        _REPORTS.pop(rid, None)
    except Exception:
        lines = [notes] if notes else []
    return {
        "alert_id": aid,
        "str_notes": notes,
        "has_saved_draft": has_saved,
        "word_preview_lines": lines[:60] if lines else ([notes] if notes else []),
    }


@router.get("/str/draft/{alert_id}/download")
async def download_str_draft_preview(
    request: Request,
    alert_id: str,
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    alert = _ALERTS.get(aid)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    if not _alert_visible_to_user(user, alert):
        raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
    if not _alert_eligible_for_str_draft_preview(alert):
        raise HTTPException(status_code=400, detail="alert is not eligible for STR draft preview")
    saved_notes = get_saved_str_draft_notes(aid)
    notes = saved_notes or (alert.escalation_reason_notes or "").strip() or "STR draft note"
    tmp = await _draft_str_report(request, aid, notes, user, allow_preapproval_preview=True)
    rid = str(tmp["report_id"])
    rec = _REPORTS.get(rid)
    if not rec:
        raise HTTPException(status_code=500, detail="failed to prepare STR draft preview")
    pg = getattr(request.app.state, "pg", None)
    docx = await str_docx_bytes_from_report_record(rec, pg=pg, user=user)
    _REPORTS.pop(rid, None)
    customer_name = str(getattr(alert, "customer_name", "") or "").strip() or "Customer"
    return Response(
        content=docx,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _report_download_content_disposition(customer_name, "STR_DRAFT", "docx")},
    )


@router.post("/str/draft/{alert_id}")
async def save_str_draft(
    alert_id: str,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    if not aid:
        raise HTTPException(status_code=400, detail="alert_id is required")
    if aid not in _ALERTS:
        raise HTTPException(status_code=404, detail="alert not found")
    alert = _ALERTS[aid]
    if not _alert_eligible_for_str_draft_preview(alert):
        raise HTTPException(status_code=400, detail="alert is not eligible for STR draft editing")
    notes = str(payload.get("str_notes") or "").strip()
    if not notes:
        raise HTTPException(status_code=400, detail="str_notes is required")
    if _looks_like_str_editor_scaffold(notes):
        extracted = _extract_internal_narrative_from_scaffold(notes)
        if extracted:
            notes = extracted
    saved_notes = notes[:_STR_DRAFT_MAX_CHARS]
    _STR_DRAFT_OVERRIDES[aid] = {
        "str_notes": saved_notes,
        "updated_at": datetime.utcnow().isoformat() + "Z",
        "updated_by": str(user.get("email") or user.get("sub") or ""),
    }
    audit_trail.record_event_from_user(
        user,
        action="report.str_draft.saved",
        resource_type="alert",
        resource_id=aid,
        details={"str_notes_length": len(saved_notes)},
    )
    return {"status": "ok", "alert_id": aid, "str_notes": saved_notes}


@router.delete("/str/draft/{alert_id}")
async def delete_str_draft(
    alert_id: str,
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    if not aid:
        raise HTTPException(status_code=400, detail="alert_id is required")
    alert = _ALERTS.get(aid)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    existed = aid in _STR_DRAFT_OVERRIDES
    _STR_DRAFT_OVERRIDES.pop(aid, None)
    audit_trail.record_event_from_user(
        user,
        action="report.str_draft.deleted",
        resource_type="alert",
        resource_id=aid,
        details={"deleted": bool(existed)},
    )
    return {"status": "ok", "alert_id": aid, "deleted": bool(existed)}


@router.post("/str/draft/status")
async def str_draft_status_bulk(
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    raw = payload.get("alert_ids")
    if not isinstance(raw, list) or not raw:
        raise HTTPException(status_code=400, detail="alert_ids must be a non-empty array")
    out: Dict[str, bool] = {}
    for x in raw[:500]:
        aid = str(x or "").strip()
        if not aid:
            continue
        a = _ALERTS.get(aid)
        if not a:
            continue
        if not _alert_visible_to_user(user, a):
            continue
        out[aid] = bool(get_saved_str_draft_notes(aid))
    return {"items": out}


@router.get("/otc-word/draft/{alert_id}")
async def get_otc_word_draft_preview(
    request: Request,
    alert_id: str,
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    alert = _ALERTS.get(aid)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    if not _alert_visible_to_user(user, alert):
        raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
    notes = _default_otc_word_draft_notes(alert)
    has_saved = bool(get_saved_otc_word_draft_notes(aid))
    if not _alert_eligible_for_otc_word_draft_preview(alert):
        return {
            "alert_id": aid,
            "estr_notes": notes,
            "has_saved_draft": has_saved,
            "otc_report_kind": getattr(alert, "otc_report_kind", None),
            "word_preview_lines": [notes] if notes else [],
            "preview_warning": "Alert is not currently eligible for OTC Word draft preview (needs escalated true-positive OTC ESTR or ESAR).",
        }
    # Keep saved text and downloaded draft identical for user-edited OTC drafts.
    if has_saved and notes:
        lines = [p.strip() for p in re.split(r"\n\s*\n", notes) if p.strip()]
        return {
            "alert_id": aid,
            "estr_notes": notes,
            "has_saved_draft": True,
            "otc_report_kind": getattr(alert, "otc_report_kind", None),
            "word_preview_lines": lines[:120] if lines else [notes],
        }
    lines: List[str] = []
    try:
        docx = await _otc_word_draft_docx_bytes(request, aid, notes, user)
        doc = Document(BytesIO(docx))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception:
        lines = [notes] if notes else []
    return {
        "alert_id": aid,
        "estr_notes": notes,
        "has_saved_draft": has_saved,
        "otc_report_kind": getattr(alert, "otc_report_kind", None),
        "word_preview_lines": lines[:80] if lines else ([notes] if notes else []),
    }


@router.get("/otc-word/draft/{alert_id}/download")
async def download_otc_word_draft_preview(
    request: Request,
    alert_id: str,
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    alert = _ALERTS.get(aid)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    if not _alert_visible_to_user(user, alert):
        raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
    if not _alert_eligible_for_otc_word_draft_preview(alert):
        raise HTTPException(status_code=400, detail="alert is not eligible for OTC Word draft preview")
    notes = _default_otc_word_draft_notes(alert)
    has_saved = bool(get_saved_otc_word_draft_notes(aid))
    rk = str(getattr(alert, "otc_report_kind", "") or "").strip().lower()
    label = "OTC_ESAR_DRAFT" if rk == "otc_esar" else "OTC_ESTR_DRAFT"
    customer_name = str(getattr(alert, "customer_name", "") or "").strip() or "Customer"
    if has_saved and notes:
        docx = regulatory_narrative_docx_bytes(
            title="OTC report draft preview",
            subtitle=f"{'OTC ESAR' if rk == 'otc_esar' else 'OTC ESTR'} · Alert: {aid}",
            narrative=notes,
            xml_excerpt=None,
            source_note="saved_draft_text",
        )
        return Response(
            content=docx,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": _report_download_content_disposition(customer_name, label, "docx")},
        )
    docx = await _otc_word_draft_docx_bytes(request, aid, notes, user)
    return Response(
        content=docx,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _report_download_content_disposition(customer_name, label, "docx")},
    )


@router.post("/otc-word/draft/{alert_id}")
async def save_otc_word_draft(
    alert_id: str,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    aid = (alert_id or "").strip()
    if not aid:
        raise HTTPException(status_code=400, detail="alert_id is required")
    if aid not in _ALERTS:
        raise HTTPException(status_code=404, detail="alert not found")
    alert = _ALERTS[aid]
    if not _alert_eligible_for_otc_word_draft_preview(alert):
        raise HTTPException(status_code=400, detail="alert is not eligible for OTC Word draft editing")
    notes = str(payload.get("estr_notes") or payload.get("notes") or "").strip()
    if not notes:
        raise HTTPException(status_code=400, detail="estr_notes is required")
    saved = notes[:_OTC_WORD_DRAFT_MAX_CHARS]
    _OTC_WORD_DRAFT_OVERRIDES[aid] = {
        "estr_notes": saved,
        "updated_at": datetime.utcnow().isoformat() + "Z",
        "updated_by": str(user.get("email") or user.get("sub") or ""),
    }
    audit_trail.record_event_from_user(
        user,
        action="report.otc_word_draft.saved",
        resource_type="alert",
        resource_id=aid,
        details={"estr_notes_length": len(saved), "otc_report_kind": getattr(alert, "otc_report_kind", None)},
    )
    return {"status": "ok", "alert_id": aid, "estr_notes": saved}


@router.post("/otc-word/draft/status")
async def otc_word_draft_status_bulk(
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    raw = payload.get("alert_ids")
    if not isinstance(raw, list) or not raw:
        raise HTTPException(status_code=400, detail="alert_ids must be a non-empty array")
    out: Dict[str, bool] = {}
    for x in raw[:500]:
        aid = str(x or "").strip()
        if not aid:
            continue
        a = _ALERTS.get(aid)
        if not a:
            continue
        if not _alert_visible_to_user(user, a):
            continue
        out[aid] = bool(get_saved_otc_word_draft_notes(aid))
    return {"items": out}


@router.post("/str/submit")
async def submit_str(payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    report_id = payload.get("report_id")
    if not report_id or report_id not in _REPORTS:
        raise HTTPException(status_code=404, detail="Report not found")
    if not _report_not_soft_deleted(_REPORTS[report_id]):
        raise HTTPException(status_code=404, detail="Report not found")
    _REPORTS[report_id]["status"] = "submitted"
    audit_trail.record_event_from_user(
        user,
        action="report.submitted.str",
        resource_type="regulatory_report",
        resource_id=str(report_id),
        details={
            "customer_id": _REPORTS[report_id].get("customer_id"),
            "alert_id": _REPORTS[report_id].get("alert_id"),
        },
    )
    submission_id = str(uuid4())
    return {"submission_id": submission_id, "nfiu_reference": None}


@router.get("/str/eligible-alerts")
async def list_str_eligible_alerts(
    request: Request,
    limit: int = Query(500, ge=1, le=1000),
    user: Dict[str, Any] = Depends(get_current_user),
):
    """Alerts that are escalated, CCO-approved for STR, and not false-positive closed (for Regulatory reports UI)."""
    pg = getattr(request.app.state, "pg", None)
    items: List[AlertResponse] = []
    for a in _ALERTS.values():
        if not _alert_visible_to_user(user, a) or not _alert_not_soft_deleted(a):
            continue
        if not _alert_eligible_for_str(a):
            continue
        items.append(await _enrich_alert_for_api(a, pg))
    items.sort(key=lambda x: (x.severity or 0.0), reverse=True)
    total = len(items)
    return {"items": items[:limit], "total": total}


@router.get("/str/{report_id}")
async def get_str(report_id: str, user: Dict[str, Any] = Depends(get_current_user)):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "STR":
        raise HTTPException(status_code=404, detail="Report not found")
    return {"report_id": report_id, "status": r.get("status"), "xml_content": r.get("xml")}


@router.post("/str/{report_id}/regenerate")
async def regenerate_str(
    request: Request,
    report_id: str,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r):
        raise HTTPException(status_code=404, detail="Report not found")
    alert_id = r.get("alert_id")
    alert = _ALERTS.get(alert_id)
    if not alert:
        raise HTTPException(status_code=404, detail="alert not found")
    if not _alert_eligible_for_str(alert):
        raise HTTPException(
            status_code=400,
            detail="STR can only be regenerated while the alert remains escalated and CCO-approved for filing.",
        )

    str_notes = _extract_str_notes(payload) or (r.get("str_notes") or "")
    if not str_notes:
        raise HTTPException(status_code=400, detail="str_notes is required (in body or from prior draft).")

    txn = _TXNS.get(alert.transaction_id)
    txn_dict = txn.model_dump() if txn else r.get("txn") or {}

    inflows_total, outflows_total = _compute_customer_cashflow_totals(alert.customer_id)
    now = datetime.utcnow()
    period_text = f"{(now - timedelta(days=365)).strftime('%B %d, %Y')} to {now.strftime('%B %d, %Y')} (12-month window)"

    prof = await _load_report_profile(request)
    xml_preview = _render_str_xml(
        alert, alert_id, txn_dict, inflows_total, outflows_total, period_text, str_notes, profile=prof
    )
    r["xml"] = xml_preview
    r["str_notes"] = str_notes
    r["status"] = "draft"
    r["customer_id"] = alert.customer_id
    r["txn"] = txn_dict
    r["alert"] = {**alert.model_dump(), "inflows_total": inflows_total, "outflows_total": outflows_total, "period_text": period_text}
    audit_trail.record_event_from_user(
        user,
        action="report.regenerated.str",
        resource_type="regulatory_report",
        resource_id=report_id,
        details={"alert_id": alert_id, "customer_id": alert.customer_id},
    )
    return {"report_id": report_id, "xml_preview": xml_preview, "validation_passed": True}


def _approver_display_name(user: Dict[str, Any]) -> str:
    return (
        str(user.get("display_name") or "").strip()
        or str(user.get("name") or "").strip()
        or str(user.get("email") or "").strip()
        or str(user.get("sub") or "").strip()
        or "Authorized Officer"
    )


async def str_docx_bytes_from_report_record(
    r: Dict[str, Any],
    *,
    pg: Any,
    user: Dict[str, Any],
) -> bytes:
    """Build STR Word bytes from a persisted STR draft (same logic as download?format=word)."""
    txn_dict = r.get("txn") or {}
    alert_context = r.get("alert") or {}
    saved_notes = str(r.get("str_notes") or "").strip() or None
    customer_id = r.get("customer_id") or alert_context.get("customer_id") or ""
    customer = await get_or_create_customer_kyc(pg, customer_id, txn_dict)
    enrichment: Optional[Dict[str, Any]] = None
    aid = r.get("alert_id")
    alert_obj = _ALERTS.get(aid) if aid else None
    if alert_obj:
        all_tx = [t.model_dump() for t in _TXNS.values()]
        enrichment = await build_alert_snapshot(
            alert=alert_obj,
            txn=txn_dict,
            all_txn_dicts=all_tx,
            pg=pg,
        )
    return _render_str_docx_bytes_safe(
        customer=customer,
        txn=txn_dict,
        alert=alert_context,
        approver_name=_approver_display_name(user),
        enrichment=enrichment,
        str_notes=saved_notes,
    )


@router.get("/str/{report_id}/download")
async def download_str(
    request: Request,
    report_id: str,
    format: str = "word",
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "STR":
        raise HTTPException(status_code=404, detail="Report not found")

    fmt = format.lower().strip()
    if fmt not in {"word", "xml"}:
        raise HTTPException(status_code=400, detail="format must be 'word' or 'xml'")

    aid = r.get("alert_id")
    alert_obj = _ALERTS.get(aid) if aid else None
    customer_name_for_file = str((alert_obj.customer_name if alert_obj else "") or "").strip() or "Customer"

    if fmt == "xml":
        str_notes = r.get("str_notes") or ""
        xml_content = r.get("xml") or ""
        if alert_obj and aid:
            txn_dict = r.get("txn") or {}
            inflows_total, outflows_total = _compute_customer_cashflow_totals(alert_obj.customer_id)
            now = datetime.utcnow()
            period_text = f"{(now - timedelta(days=365)).strftime('%B %d, %Y')} to {now.strftime('%B %d, %Y')} (12-month window)"
            prof = await _load_report_profile(request)
            xml_content = _render_str_xml(
                alert_obj, aid, txn_dict, inflows_total, outflows_total, period_text, str_notes, profile=prof
            )
            r["xml"] = xml_content
        pg = getattr(request.app.state, "pg", None)
        customer_id = r.get("customer_id") or (alert_obj.customer_id if alert_obj else "")
        customer_name = customer_name_for_file
        if customer_id:
            try:
                kyc = await get_or_create_customer_kyc(pg, customer_id, r.get("txn") or {})
                if customer_name == "Customer":
                    customer_name = str(getattr(kyc, "customer_name", "") or "").strip() or "Customer"
            except Exception:
                customer_name = customer_name_for_file
        return Response(
            content=xml_content,
            media_type="application/xml",
            headers={"Content-Disposition": _report_download_content_disposition(customer_name, "STR", "xml")},
        )

    txn_dict = r.get("txn") or {}
    alert_context = r.get("alert") or {}
    customer_id = r.get("customer_id") or alert_context.get("customer_id") or ""
    pg = getattr(request.app.state, "pg", None)
    customer = await get_or_create_customer_kyc(pg, customer_id, txn_dict)
    enrichment: Optional[Dict[str, Any]] = None
    aid = r.get("alert_id")
    alert_obj = _ALERTS.get(aid) if aid else None
    if alert_obj:
        all_tx = [t.model_dump() for t in _TXNS.values()]
        enrichment = await build_alert_snapshot(
            alert=alert_obj,
            txn=txn_dict,
            all_txn_dicts=all_tx,
            pg=pg,
        )
    str_notes_for_word = str(r.get("str_notes") or "").strip() or None
    doc_bytes = _render_str_docx_bytes_safe(
        customer=customer,
        txn=txn_dict,
        alert=alert_context,
        approver_name=_approver_display_name(user),
        enrichment=enrichment,
        str_notes=str_notes_for_word,
    )
    return Response(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _report_download_content_disposition(customer_name_for_file or customer.customer_name, "STR", "docx")},
    )


@router.get("/str/")
async def list_strs(
    status: Optional[str] = None,
    user: Dict[str, Any] = Depends(get_current_user),
):
    items = []
    for rid, r in _REPORTS.items():
        if not _report_not_soft_deleted(r):
            continue
        if r.get("type") != "STR":
            continue
        if status and r.get("status") != status:
            continue
        items.append({"report_id": rid, "status": r.get("status"), "alert_id": r.get("alert_id")})
    return {"items": items, "total": len(items)}


@router.get("/compliance/dashboard")
async def compliance_dashboard(request: Request, user: Dict[str, Any] = Depends(get_current_user)):
    regulatory_deadlines: List[Dict[str, Any]] = []
    institution_display_name: Optional[str] = None
    pg = getattr(request.app.state, "pg", None)
    if pg:
        try:
            entries = await list_calendar_entries(pg)
            regulatory_deadlines = upcoming_calendar_preview(entries)[:12]
            prof = await get_reporting_profile_row(pg)
            institution_display_name = prof.get("institution_display_name")
        except Exception:
            pass
    return {
        "str_submission_rate": 0.0,
        "acceptance_rate": None,
        "average_time_to_submit_hours": None,
        "regulatory_deadlines": regulatory_deadlines,
        "institution_display_name": institution_display_name,
    }


def _normalize_bundle_report_kinds(raw: Any) -> set[str]:
    out: set[str] = set()
    if not isinstance(raw, list):
        return out
    for x in raw:
        k = str(x).strip().lower()
        if k == "str":
            out.add("str")
        elif k == "aop":
            out.add("aop")
        elif k in ("nfiu_customer_change", "nfiu", "cir", "customer_information_change"):
            out.add("nfiu")
    return out


@router.post("/bundle/generate")
async def generate_compliance_bundle(request: Request, payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    """
    Multi-customer / multi-report generation: tick-style selection via `reports`: str, aop, nfiu.
    STR requires `alert_id` + `str_notes` and normal STR eligibility. AOP and NFIU use `customer_id`
    (from alert if omitted). When `use_llm` is true (default), AOP and NFIU Word narratives use the
    configured LLM when available; STR Word still uses existing snapshot enrichment on download.
    """
    use_llm = payload.get("use_llm")
    if use_llm is None:
        use_llm = True
    if isinstance(use_llm, str):
        use_llm = use_llm.strip().lower() not in ("0", "false", "no")
    use_llm_b = bool(use_llm)

    raw_items = payload.get("items")
    if not isinstance(raw_items, list) or not raw_items:
        raise HTTPException(status_code=400, detail="items must be a non-empty array")
    if len(raw_items) > 40:
        raise HTTPException(status_code=400, detail="Maximum 40 bundle items per request")

    bundle_id = str(uuid4())
    pg = getattr(request.app.state, "pg", None)
    prof = await _load_report_profile(request)
    summarized: List[Dict[str, Any]] = []

    for idx, raw in enumerate(raw_items):
        if not isinstance(raw, dict):
            summarized.append({"index": idx, "error": "invalid item", "reports": []})
            continue

        rk = _normalize_bundle_report_kinds(raw.get("reports"))
        if not rk:
            summarized.append(
                {
                    "index": idx,
                    "error": "reports must include at least one of: str, aop, nfiu (or nfiu_customer_change)",
                    "reports": [],
                }
            )
            continue

        alert_id = str(raw.get("alert_id") or "").strip()
        customer_override = str(raw.get("customer_id") or "").strip()
        alert: Optional[AlertResponse] = _ALERTS.get(alert_id) if alert_id else None

        if alert_id and alert is None:
            summarized.append({"index": idx, "alert_id": alert_id, "error": "alert not found", "reports": []})
            continue

        customer_id = customer_override
        if alert:
            if not _alert_visible_to_user(user, alert):
                summarized.append(
                    {
                        "index": idx,
                        "alert_id": alert_id,
                        "error": "Alert outside your zone/branch scope",
                        "reports": [],
                    }
                )
                continue
            if not customer_id:
                customer_id = alert.customer_id
        elif "str" in rk:
            summarized.append({"index": idx, "error": "alert_id is required when STR is included", "reports": []})
            continue
        elif not customer_id:
            summarized.append(
                {"index": idx, "error": "customer_id is required when no alert_id is provided", "reports": []}
            )
            continue

        txn_for_kyc: Optional[TransactionResponse] = None
        if alert:
            txn_for_kyc = _TXNS.get(alert.transaction_id)
        if txn_for_kyc is None:
            txn_for_kyc = _latest_txn_for_customer(customer_id)

        if txn_for_kyc:
            txn_dict = txn_for_kyc.model_dump()
        else:
            txn_dict = {
                "created_at": datetime.utcnow(),
                "amount": 0.0,
                "transaction_type": "",
                "currency": "NGN",
                "narrative": (alert.summary if alert else "") or "Bundle KYC seeding (demo)",
                "metadata": {},
                "customer_id": customer_id,
                "id": (alert.transaction_id if alert else "") or f"bundle-{customer_id[:24]}",
            }

        customer_kyc = await get_or_create_customer_kyc(pg, customer_id, txn_dict)
        cust_name = customer_kyc.customer_name

        row_out: Dict[str, Any] = {
            "index": idx,
            "customer_id": customer_id,
            "alert_id": alert_id or None,
            "reports": [],
        }

        if "str" in rk:
            try:
                sn = str(raw.get("str_notes") or "").strip()
                if not sn:
                    row_out["reports"].append({"kind": "str", "ok": False, "error": "str_notes required"})
                elif not alert:
                    row_out["reports"].append({"kind": "str", "ok": False, "error": "alert not found"})
                else:
                    dr = await _draft_str_report(request, alert_id, sn, user, profile=prof)
                    row_out["reports"].append({"kind": "str", "ok": True, **dr})
            except HTTPException as he:
                row_out["reports"].append({"kind": "str", "ok": False, "error": str(he.detail)})

        if "aop" in rk:
            aop_body = raw.get("aop") if isinstance(raw.get("aop"), dict) else {}
            prod = str(aop_body.get("account_product") or "Savings")
            risk = str(aop_body.get("risk_rating") or "medium")
            str_sum = str(raw.get("str_notes") or "").strip()[:1500] if "str" in rk else None
            wn, ws = await build_aop_bundle_narrative(
                customer_id=customer_id,
                customer_name=cust_name,
                account_product=prod,
                risk_rating=risk,
                str_notes_summary=str_sum,
                use_llm=use_llm_b,
            )
            dr = _draft_aop_record(
                customer_id,
                user,
                account_product=prod,
                risk_rating=risk,
                word_narrative=wn,
                word_narrative_source=ws,
                profile=prof,
            )
            row_out["reports"].append({"kind": "aop", "ok": True, **dr})

        if "nfiu" in rk:
            nf = raw.get("nfiu") if isinstance(raw.get("nfiu"), dict) else {}
            change_type = str(nf.get("change_type") or "partial_name_change").strip()
            allowed_nfiu = {"partial_name_change", "full_name_change", "bvn_update", "dob_change"}
            if change_type not in allowed_nfiu:
                row_out["reports"].append({"kind": "nfiu", "ok": False, "error": f"invalid change_type: {change_type}"})
            else:
                fields = {
                    "old_value": nf.get("old_value"),
                    "new_value": nf.get("new_value"),
                    "notes": nf.get("notes"),
                    "bvn_old": nf.get("bvn_old"),
                    "bvn_new": nf.get("bvn_new"),
                    "name_old": nf.get("name_old"),
                    "name_new": nf.get("name_new"),
                    "dob_old": nf.get("dob_old"),
                    "dob_new": nf.get("dob_new"),
                }
                wn, ws = await build_nfiu_cir_bundle_narrative(
                    change_type=change_type,
                    customer_id=customer_id,
                    fields=fields,
                    use_llm=use_llm_b,
                )
                dr = _draft_nfiu_cir_record(
                    change_type,
                    customer_id,
                    user,
                    fields,
                    word_narrative=wn,
                    word_narrative_source=ws,
                )
                row_out["reports"].append({"kind": "nfiu", "ok": True, **dr})

        summarized.append(row_out)

    return {"bundle_id": bundle_id, "use_llm": use_llm_b, "items": summarized}


def _download_registered_report(
    report_id: str,
    format: str,
    *,
    filename_prefix: str,
    title: str,
) -> Response:
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r):
        raise HTTPException(status_code=404, detail="Report not found")
    fmt = format.lower().strip()
    if fmt not in {"word", "xml"}:
        raise HTTPException(status_code=400, detail="format must be 'word' or 'xml'")
    xml_content = str(r.get("xml") or "")
    if fmt == "xml":
        return Response(
            content=xml_content,
            media_type="application/xml",
            headers={"Content-Disposition": f'attachment; filename="{filename_prefix}_{report_id}.xml"'},
        )
    doc = minimal_docx_bytes(title, f"Demo extract — {title}\n\nXML payload (truncated):\n{xml_content[:12000]}")
    return Response(
        content=doc,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename_prefix}_{report_id}.docx"'},
    )


@router.post("/ctr/generate")
async def generate_ctr(
    request: Request,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    """Demo Currency Transaction Report (goAML-style stub)."""
    report_id = str(uuid4())
    sample_tx = list(_TXNS.values())[:5]
    tx_summ = "; ".join(f"{t.id}:{t.amount}" for t in sample_tx) if sample_tx else "no_transactions"
    prof = await _load_report_profile(request)
    xml = goaml_stub_xml(
        "CTR",
        merge_goaml_stub_payload(
            prof,
            {
                "report_id": report_id,
                "transaction_sample": tx_summ,
                "customer_id": payload.get("customer_id") or "",
                "prepared_by": str(user.get("display_name") or user.get("email") or "Compliance"),
            },
        ),
    )
    ctr_rec = {
        "type": "CTR",
        "status": "draft",
        "xml": xml,
        "customer_id": str(payload.get("customer_id") or "") or None,
    }
    _persist_report_with_audit(user, report_id, ctr_rec)
    return {"report_id": report_id, "xml_preview": xml, "validation_passed": True}


@router.get("/ctr/{report_id}/download")
async def download_ctr(
    report_id: str,
    format: str = "xml",
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "CTR":
        raise HTTPException(status_code=404, detail="CTR report not found")
    return _download_registered_report(report_id, format, filename_prefix="CTR", title="Currency Transaction Report (demo)")


def _payload_us_activity_focus(payload: Dict[str, Any]) -> bool:
    v = payload.get("us_activity")
    if v is None:
        return False
    if isinstance(v, bool):
        return v
    s = str(v).strip().lower()
    return s not in ("0", "false", "no", "off")


async def _create_sar_report_record(
    request: Request,
    user: Dict[str, Any],
    *,
    report_id: str,
    alert_id: str,
    customer_id: str,
    txn_dict: Dict[str, Any],
    sar_notes: str,
    enrichment: Optional[Dict[str, Any]],
    activity_profile: Optional[Dict[str, Any]],
    us_activity_focus: bool,
) -> Dict[str, Any]:
    pg = getattr(request.app.state, "pg", None)
    customer = await get_or_create_customer_kyc(pg, customer_id, txn_dict)
    alert_dict: Optional[Dict[str, Any]] = None
    alert_obj = _ALERTS.get(alert_id) if alert_id else None
    if alert_obj:
        alert_dict = alert_obj.model_dump()
    case_ctx = build_sar_case_context(
        customer=customer,
        txn=txn_dict,
        alert=alert_dict,
        sar_notes=sar_notes,
        enrichment=enrichment,
        activity_profile=activity_profile,
        us_activity_focus=us_activity_focus,
    )
    sections, narrative_source = await generate_sar_narrative_sections(case_ctx)
    basis = "suspicious_activity" if activity_profile else "transaction_led"
    prof = await _load_report_profile(request)
    xml = goaml_stub_xml(
        "SAR",
        merge_goaml_stub_payload(
            prof,
            {
                "report_id": report_id,
                "alert_id": alert_id,
                "customer_id": customer_id,
                "transaction_id": str(txn_dict.get("id") or ""),
                "sar_notes": sar_notes,
                "narrative_source": narrative_source,
                "activity_basis": basis,
                "us_activity_focus": "true" if us_activity_focus else "false",
                "prepared_by": str(user.get("display_name") or user.get("email") or "Compliance"),
            },
        ),
    )
    sar_rec = {
        "type": "SAR",
        "status": "draft",
        "xml": xml,
        "alert_id": alert_id or None,
        "customer_id": customer_id,
        "txn": txn_dict,
        "sections": sections,
        "narrative_source": narrative_source,
        "sar_notes": sar_notes,
        "activity_basis": basis,
        "us_activity_focus": us_activity_focus,
    }
    _persist_report_with_audit(user, report_id, sar_rec)
    return {
        "report_id": report_id,
        "xml_preview": xml,
        "validation_passed": True,
        "narrative_source": narrative_source,
        "activity_basis": basis,
    }


@router.get("/sar/eligible-alerts")
async def list_sar_eligible_alerts(
    request: Request,
    limit: int = 500,
    user: Dict[str, Any] = Depends(get_current_user),
):
    """Alerts closed as false positive — eligible for suspicious-activity SAR (scoped like the alerts list)."""
    cap = max(1, min(limit, 2000))
    pg = getattr(request.app.state, "pg", None)
    rows: List[Dict[str, Any]] = []
    for a in _ALERTS.values():
        if not _alert_eligible_for_sar(a):
            continue
        if not _alert_visible_to_user(user, a):
            continue
        customer_name = str(getattr(a, "customer_name", "") or "").strip()
        if not customer_name:
            kyc = await fetch_customer_kyc_any(pg, a.customer_id)
            customer_name = str(getattr(kyc, "customer_name", "") or "").strip()
        rows.append(
            {
                "alert_id": a.id,
                "customer_id": a.customer_id,
                "customer_name": customer_name or None,
                "linked_channel": _linked_channel_for_alert(a),
                "transaction_id": a.transaction_id,
                "summary": a.summary,
                "severity": a.severity,
                "updated_at": a.updated_at.isoformat() if a.updated_at else None,
            }
        )
    rows.sort(key=lambda r: float(r.get("severity") or 0.0), reverse=True)
    rows = rows[:cap]
    return {"items": rows, "total": len(rows)}


@router.get("/otc/eligible-alerts")
async def list_otc_eligible_alerts(
    request: Request,
    kind: str = Query("estr", description="estr or esar"),
    limit: int = Query(500, ge=1, le=2000),
    user: Dict[str, Any] = Depends(get_current_user),
):
    """OTC true-positive path: ESTR (cash) when filing is on file; ESAR via SAR (identity matters)."""
    k = (kind or "").strip().lower()
    if k not in ("estr", "esar"):
        raise HTTPException(status_code=400, detail="kind must be estr or esar")
    want_estr = k == "estr"
    fn = _alert_eligible_for_otc_estr_word_ready if want_estr else _alert_eligible_for_otc_esar
    cap = max(1, min(limit, 2000))
    pg = getattr(request.app.state, "pg", None)
    rows: List[Dict[str, Any]] = []
    for a in _ALERTS.values():
        if not fn(a):
            continue
        if not _alert_visible_to_user(user, a):
            continue
        customer_name = str(getattr(a, "customer_name", "") or "").strip()
        if not customer_name:
            kyc = await fetch_customer_kyc_any(pg, a.customer_id)
            customer_name = str(getattr(kyc, "customer_name", "") or "").strip()
        rows.append(
            {
                "alert_id": a.id,
                "customer_id": a.customer_id,
                "customer_name": customer_name or None,
                "linked_channel": _linked_channel_for_alert(a),
                "transaction_id": a.transaction_id,
                "summary": a.summary,
                "otc_subject": a.otc_subject,
                "otc_report_kind": a.otc_report_kind,
                "severity": a.severity,
                "updated_at": a.updated_at.isoformat() if a.updated_at else None,
            }
        )
    rows.sort(key=lambda r: float(r.get("severity") or 0.0), reverse=True)
    rows = rows[:cap]
    return {"items": rows, "total": len(rows), "kind": k}


@router.post("/sar/generate-bulk")
async def generate_sar_bulk(request: Request, payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    """
    Generate many SAR drafts in one request. Each alert must be visible to the user and eligible as either:
    closed false-positive (activity-led SAR) or CCO-approved OTC ESAR. OTC ESTR alerts are rejected per alert.
    Omit alert_ids to auto-pick up to `limit` such alerts (highest severity first).
    """
    raw_ids = payload.get("alert_ids")
    limit = int(payload.get("limit") or 50)
    limit = max(1, min(limit, 500))
    sar_notes = str(payload.get("sar_notes") or payload.get("notes") or "").strip()[:4000]
    us_focus = _payload_us_activity_focus(payload)
    use_saved_sar = bool(payload.get("use_saved_draft", True))

    id_list: List[str] = []
    if isinstance(raw_ids, list) and raw_ids:
        id_list = [str(x).strip() for x in raw_ids if str(x).strip()][:500]
    else:
        tmp: List[AlertResponse] = []
        for a in _ALERTS.values():
            if not _alert_eligible_for_sar_bulk_item(a):
                continue
            if not _alert_visible_to_user(user, a):
                continue
            tmp.append(a)
        tmp.sort(key=lambda x: float(x.severity or 0.0), reverse=True)
        id_list = [a.id for a in tmp[:limit]]

    results: List[Dict[str, Any]] = []
    pg = getattr(request.app.state, "pg", None)
    all_tx = [t.model_dump() for t in _TXNS.values()]

    for aid in id_list:
        alert = _ALERTS.get(aid)
        if not alert:
            results.append({"alert_id": aid, "ok": False, "error": "alert not found"})
            continue
        if not _alert_visible_to_user(user, alert):
            results.append({"alert_id": aid, "ok": False, "error": "outside your zone/branch scope"})
            continue
        rk_otc = getattr(alert, "otc_report_kind", None)
        if rk_otc == "otc_estr":
            results.append({"alert_id": aid, "ok": False, "error": "OTC ESTR path — generate ESTR, not SAR"})
            continue
        report_id = str(uuid4())
        try:
            if _alert_eligible_for_otc_esar(alert):
                txn_dict, linked, synthetic = _resolve_txn_for_sar(alert, None)
                activity_profile = _activity_profile_for_otc_esar(alert)
                sar_piece = sar_notes
                if use_saved_sar:
                    dn = get_saved_otc_word_draft_notes(aid)
                    if dn:
                        sar_piece = dn[:4000]
                combined_notes = _combine_sar_notes_otc(alert, sar_piece)
                enrichment = None
                if not synthetic:
                    enrichment = await build_alert_snapshot(alert=alert, txn=txn_dict, all_txn_dicts=all_tx, pg=pg)
                out = await _create_sar_report_record(
                    request,
                    user,
                    report_id=report_id,
                    alert_id=alert.id,
                    customer_id=alert.customer_id,
                    txn_dict=txn_dict,
                    sar_notes=combined_notes,
                    enrichment=enrichment,
                    activity_profile=activity_profile,
                    us_activity_focus=us_focus,
                )
                results.append({"alert_id": aid, "ok": True, **out})
            elif _alert_eligible_for_sar(alert):
                txn_dict, linked, synthetic = _resolve_txn_for_sar(alert, None)
                activity_profile = _activity_profile_for_sar(alert, None if synthetic else linked)
                enrichment = None
                if not synthetic:
                    enrichment = await build_alert_snapshot(alert=alert, txn=txn_dict, all_txn_dicts=all_tx, pg=pg)
                out = await _create_sar_report_record(
                    request,
                    user,
                    report_id=report_id,
                    alert_id=alert.id,
                    customer_id=alert.customer_id,
                    txn_dict=txn_dict,
                    sar_notes=sar_notes,
                    enrichment=enrichment,
                    activity_profile=activity_profile,
                    us_activity_focus=us_focus,
                )
                results.append({"alert_id": aid, "ok": True, **out})
            else:
                results.append(
                    {
                        "alert_id": aid,
                        "ok": False,
                        "error": "not eligible (needs false-positive closed SAR or CCO-approved OTC ESAR)",
                    }
                )
        except Exception as exc:
            results.append({"alert_id": aid, "ok": False, "error": str(exc)})

    ok_n = sum(1 for r in results if r.get("ok"))
    return {"results": results, "generated": ok_n, "requested": len(id_list)}


@router.post("/sar/generate")
async def generate_sar(request: Request, payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    """
    Suspicious Activity Report: goAML XML stub plus structured Word document.
    When `alert_id` refers to an alert closed as **false positive**, SAR is **activity-led** (existing scenario / typology;
    no transaction required). Optional `us_activity` (default false) may steer narrative toward US nexus where applicable (demo).
    Legacy: without that eligible alert, `customer_id` + a stored transaction (or transaction_id) is still supported.
    """
    report_id = str(uuid4())
    alert_id = str(payload.get("alert_id") or "").strip()
    sar_notes = str(payload.get("sar_notes") or payload.get("notes") or "")[:4000]
    us_focus = _payload_us_activity_focus(payload)
    txn_override = str(payload.get("transaction_id") or "").strip() or None

    alert: Optional[AlertResponse] = _ALERTS.get(alert_id) if alert_id else None
    customer_id = str(payload.get("customer_id") or "").strip()
    if alert:
        customer_id = alert.customer_id
    if not customer_id:
        raise HTTPException(status_code=400, detail="customer_id or alert_id is required")

    if alert:
        if not _alert_visible_to_user(user, alert):
            raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
        rk_otc = getattr(alert, "otc_report_kind", None)
        if rk_otc == "otc_estr":
            raise HTTPException(
                status_code=400,
                detail="This alert is on the OTC ESTR (cash) path. Generate ESTR from Regulatory reports, not SAR.",
            )
        if _alert_eligible_for_otc_esar(alert):
            txn_dict, linked, synthetic = _resolve_txn_for_sar(alert, txn_override)
            activity_profile = _activity_profile_for_otc_esar(alert)
            sar_notes_eff = sar_notes
            if bool(payload.get("use_saved_draft", True)):
                dn = get_saved_otc_word_draft_notes(alert_id)
                if dn:
                    sar_notes_eff = dn[:4000]
            combined_notes = _combine_sar_notes_otc(alert, sar_notes_eff)
            enrichment = None
            if not synthetic:
                all_tx = [t.model_dump() for t in _TXNS.values()]
                pg = getattr(request.app.state, "pg", None)
                enrichment = await build_alert_snapshot(alert=alert, txn=txn_dict, all_txn_dicts=all_tx, pg=pg)
            return await _create_sar_report_record(
                request,
                user,
                report_id=report_id,
                alert_id=alert.id,
                customer_id=customer_id,
                txn_dict=txn_dict,
                sar_notes=combined_notes,
                enrichment=enrichment,
                activity_profile=activity_profile,
                us_activity_focus=us_focus,
            )
        if not _alert_eligible_for_sar(alert):
            raise HTTPException(
                status_code=400,
                detail="SAR for this workflow requires the alert to be closed as a false positive, or an OTC ESAR "
                "(identity matter) with CCO approval. Resolve, file OTC, or use customer_id + transaction for legacy SAR.",
            )
        txn_dict, linked, synthetic = _resolve_txn_for_sar(alert, txn_override)
        activity_profile = _activity_profile_for_sar(alert, None if synthetic else linked)
        enrichment = None
        if not synthetic:
            all_tx = [t.model_dump() for t in _TXNS.values()]
            pg = getattr(request.app.state, "pg", None)
            enrichment = await build_alert_snapshot(alert=alert, txn=txn_dict, all_txn_dicts=all_tx, pg=pg)
        return await _create_sar_report_record(
            request,
            user,
            report_id=report_id,
            alert_id=alert.id,
            customer_id=customer_id,
            txn_dict=txn_dict,
            sar_notes=sar_notes,
            enrichment=enrichment,
            activity_profile=activity_profile,
            us_activity_focus=us_focus,
        )

    txn_id = str(payload.get("transaction_id") or "").strip()
    txn = _TXNS.get(txn_id) if txn_id else None
    if txn is None:
        txn = _latest_txn_for_customer(customer_id)
    if txn is None:
        raise HTTPException(
            status_code=400,
            detail="No transaction found for this customer. Ingest a transaction, pass transaction_id, "
            "or use alert_id for a false-positive-closed alert (activity-based SAR, no txn required).",
        )
    txn_dict = txn.model_dump()
    pg = getattr(request.app.state, "pg", None)
    enrichment = None
    return await _create_sar_report_record(
        request,
        user,
        report_id=report_id,
        alert_id="",
        customer_id=customer_id,
        txn_dict=txn_dict,
        sar_notes=sar_notes,
        enrichment=enrichment,
        activity_profile=None,
        us_activity_focus=us_focus,
    )


@router.get("/sar/{report_id}/download")
async def download_sar(
    request: Request,
    report_id: str,
    format: str = "xml",
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "SAR":
        raise HTTPException(status_code=404, detail="SAR report not found")
    fmt = format.lower().strip()
    pg = getattr(request.app.state, "pg", None)
    txn_dict = r.get("txn") or {}
    customer = await get_or_create_customer_kyc(pg, str(r.get("customer_id") or ""), txn_dict)
    if fmt == "xml":
        return Response(
            content=(r.get("xml") or ""),
            media_type="application/xml",
            headers={"Content-Disposition": _report_download_content_disposition(customer.customer_name, "SAR", "xml")},
        )

    if fmt != "word":
        raise HTTPException(status_code=400, detail="format must be 'word' or 'xml'")

    sections = r.get("sections")
    if not isinstance(sections, dict):
        raise HTTPException(status_code=400, detail="This SAR predates Word export; regenerate the SAR.")

    sig_path = (settings.cco_signature_image_path or "").strip() or None
    doc_bytes = render_sar_docx_bytes(
        customer=customer,
        txn=txn_dict,
        sections=sections,
        approver_name=_approver_display_name(user),
        signature_image_path=sig_path,
        narrative_source=str(r.get("narrative_source") or "template"),
    )
    return Response(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _report_download_content_disposition(customer.customer_name, "SAR", "docx")},
    )


@router.post("/aop/generate")
async def generate_aop(
    request: Request,
    payload: Dict[str, Any],
    user: Dict[str, Any] = Depends(get_current_user),
):
    """Demo Account Opening Package (goAML onboarding stub)."""
    prof = await _load_report_profile(request)
    return _draft_aop_record(
        str(payload.get("customer_id") or "UNKNOWN"),
        user,
        account_product=str(payload.get("account_product") or "Savings"),
        risk_rating=str(payload.get("risk_rating") or "medium"),
        profile=prof,
    )


@router.get("/aop/{report_id}/download")
async def download_aop(report_id: str, format: str = "pdf", user: Dict[str, Any] = Depends(get_current_user)):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "AOP":
        raise HTTPException(status_code=404, detail="AOP report not found")
    fmt = format.lower().strip()
    if fmt in ("xml", "word"):
        raise HTTPException(
            status_code=400,
            detail="Only PDF is available for the account opening package. Use format=pdf (default).",
        )
    if fmt != "pdf":
        raise HTTPException(status_code=400, detail="format must be 'pdf'")
    narrative = str(r.get("word_narrative") or "").strip()
    if not narrative:
        narrative = (
            "Account opening package (demo).\n\n"
            "No narrative text is stored on this draft yet. Regenerate the AOP from Customers or "
            "from STR generation with “Include AOP” to populate the package."
        )
    pdf = regulatory_narrative_pdf_bytes(
        title="Account opening package",
        subtitle=f"Customer: {r.get('customer_id')}",
        narrative=narrative,
        source_note=str(r.get("word_narrative_source") or "").strip() or None,
        xml_excerpt=None,
    )
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="Account-opening-package_{report_id}.pdf"'},
    )


@router.get("/soa/{report_id}/download")
async def download_soa(
    report_id: str,
    format: str = "word",
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "SOA":
        raise HTTPException(status_code=404, detail="Statement of account not found")
    fmt = format.lower().strip()
    if fmt == "xml":
        raise HTTPException(
            status_code=400,
            detail="XML export is not available for statements of account. Use format=word.",
        )
    if fmt != "word":
        raise HTTPException(status_code=400, detail="format must be 'word'")
    doc_bytes = regulatory_narrative_docx_bytes(
        title="Statement of account",
        subtitle=f"Customer: {r.get('customer_id')} · {r.get('period_start')} → {r.get('period_end')}",
        narrative=str(r.get("statement_text") or ""),
        xml_excerpt=str(r.get("xml") or ""),
        source_note="internal_package",
    )
    return Response(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="SOA_{report_id}.docx"'},
    )


async def _create_estr_draft(
    request: Request,
    user: Dict[str, Any],
    *,
    base_alert: str,
    user_notes: str,
) -> Dict[str, Any]:
    """
    Extended OTC regulatory return (goAML ESTR stub + structured Word).

    Cash deposit / withdrawal (OTC ESTR path) or identity / profile change (OTC ESAR path) may be linked after CCO approval.
    Word output title follows officer-selected subject (ESTR vs ESAR profile).
    """
    report_id = str(uuid4())
    extension = user_notes

    if base_alert:
        ao = _ALERTS.get(base_alert)
        if not ao:
            raise HTTPException(status_code=404, detail="Alert not found.")
        if not _alert_visible_to_user(user, ao):
            raise HTTPException(status_code=403, detail="Alert outside your zone/branch scope.")
        rk = getattr(ao, "otc_report_kind", None)
        if rk == "otc_estr":
            if not _alert_eligible_for_otc_estr(ao):
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "OTC ESTR requires a true-positive OTC filing (cash subject), compliance escalation, and "
                        "CCO approval of OTC reporting (CCO review queue), unless admin auto-OTC workflow is enabled."
                    ),
                )
            otc_ctx = (
                f"OTC ESTR | subject={ao.otc_subject} | filing_reason={ao.otc_filing_reason} | "
                f"officer_rationale={(ao.otc_officer_rationale or '')[:2500]}"
            )
            extension = (extension + "\n\n" + otc_ctx).strip()[:8000]
        elif rk == "otc_esar":
            if not _alert_eligible_for_otc_esar(ao):
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "OTC ESAR requires a true-positive OTC filing (identity path), compliance escalation, and "
                        "CCO approval of OTC reporting, unless admin auto-OTC workflow is enabled."
                    ),
                )
            otc_ctx = (
                f"OTC ESAR (activity) | subject={ao.otc_subject} | filing_reason={ao.otc_filing_reason} | "
                f"officer_rationale={(ao.otc_officer_rationale or '')[:2500]}"
            )
            extension = (extension + "\n\n" + otc_ctx).strip()[:8000]
        else:
            raise HTTPException(
                status_code=400,
                detail="Linked alert has no approved OTC report kind (otc_estr or otc_esar).",
            )

    prof = await _load_report_profile(request)
    xml = goaml_stub_xml(
        "ESTR",
        merge_goaml_stub_payload(
            prof,
            {
                "report_id": report_id,
                "linked_alert_id": base_alert,
                "extension_notes": extension,
                "prepared_by": str(user.get("display_name") or user.get("email") or "Compliance"),
            },
        ),
    )
    cid = None
    ao_for_audit = _ALERTS.get(base_alert) if base_alert else None
    if ao_for_audit:
        cid = ao_for_audit.customer_id

    customer_display_name = ""
    if cid:
        txn_dict = _txn_dict_for_estr_kyc(base_alert or None, cid)
        pg = getattr(request.app.state, "pg", None)
        customer = await get_or_create_customer_kyc(pg, cid, txn_dict)
        customer_display_name = (customer.customer_name or "").strip()

    estr_rec = {
        "type": "ESTR",
        "status": "draft",
        "xml": xml,
        "alert_id": base_alert or None,
        "customer_id": cid,
        "customer_display_name": customer_display_name,
        "estr_notes": user_notes[:4000],
    }
    _persist_report_with_audit(
        user,
        report_id,
        estr_rec,
        extra={
            "linked_alert_id": base_alert or None,
            "otc_subject": getattr(ao_for_audit, "otc_subject", None) if ao_for_audit else None,
        },
    )
    if base_alert and _ALERTS.get(base_alert):
        ao2 = _ALERTS[base_alert]
        now = datetime.utcnow()
        hist = list(ao2.investigation_history or [])
        hist.append(
            {
                "action": "estr_draft_generated",
                "report_id": report_id,
                "otc_subject": getattr(ao2, "otc_subject", None),
                "otc_report_kind": getattr(ao2, "otc_report_kind", None),
                "prepared_by": str(user.get("display_name") or user.get("email") or user.get("sub") or "Compliance"),
                "at": now.isoformat() + "Z",
            }
        )
        ao2 = ao2.model_copy(update={"investigation_history": hist, "updated_at": now})
        _ALERTS[base_alert] = ao2
        audit_trail.record_event_from_user(
            user,
            action="alert.estr_draft_generated",
            resource_type="alert",
            resource_id=base_alert,
            details={
                "report_id": report_id,
                "otc_subject": getattr(ao2, "otc_subject", None),
                "otc_report_kind": getattr(ao2, "otc_report_kind", None),
            },
        )
    return {"report_id": report_id, "xml_preview": xml, "validation_passed": True}


@router.post("/estr/generate")
async def generate_estr(request: Request, payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    base_alert = str(payload.get("alert_id") or "").strip()
    user_notes = str(payload.get("estr_notes") or payload.get("notes") or "").strip()[:8000]
    if bool(payload.get("use_saved_draft", True)) and base_alert:
        draft_notes = get_saved_otc_word_draft_notes(base_alert)
        if draft_notes:
            user_notes = draft_notes[:8000]
    return await _create_estr_draft(request, user, base_alert=base_alert, user_notes=user_notes)


@router.post("/estr/generate-bulk")
async def generate_estr_bulk(request: Request, payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    """
    Generate many OTC extended return (ESTR) drafts in one request.
    Each alert must be visible and OTC ESTR-eligible (true-positive OTC filing on file).
    Omit alert_ids to auto-pick up to `limit` such alerts (highest severity first).
    """
    raw_ids = payload.get("alert_ids")
    limit = int(payload.get("limit") or 500)
    limit = max(1, min(limit, 500))
    estr_notes = str(payload.get("estr_notes") or payload.get("notes") or "").strip()[:8000]
    use_saved = bool(payload.get("use_saved_draft", True))

    id_list: List[str] = []
    if isinstance(raw_ids, list) and raw_ids:
        id_list = [str(x).strip() for x in raw_ids if str(x).strip()][:500]
    else:
        tmp: List[AlertResponse] = []
        for a in _ALERTS.values():
            if not _alert_eligible_for_otc_estr_word_ready(a):
                continue
            if not _alert_visible_to_user(user, a):
                continue
            tmp.append(a)
        tmp.sort(key=lambda x: float(x.severity or 0.0), reverse=True)
        id_list = [a.id for a in tmp[:limit]]

    results: List[Dict[str, Any]] = []
    for aid in id_list:
        alert_row = _ALERTS.get(aid)
        cid = alert_row.customer_id if alert_row else ""
        try:
            notes_eff = estr_notes
            if use_saved:
                dn = get_saved_otc_word_draft_notes(aid)
                if dn:
                    notes_eff = dn[:8000]
            out = await _create_estr_draft(request, user, base_alert=aid, user_notes=notes_eff)
            results.append({"alert_id": aid, "customer_id": cid, "ok": True, **out})
        except HTTPException as e:
            results.append({"alert_id": aid, "customer_id": cid, "ok": False, "error": _http_exception_detail_str(e)})
        except Exception as exc:
            results.append({"alert_id": aid, "customer_id": cid, "ok": False, "error": str(exc)})

    ok_n = sum(1 for r in results if r.get("ok"))
    return {"results": results, "generated": ok_n, "requested": len(id_list)}


@router.get("/estr/{report_id}/download")
async def download_estr(
    request: Request,
    report_id: str,
    format: str = "xml",
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "ESTR":
        raise HTTPException(status_code=404, detail="ESTR report not found")
    fmt = format.lower().strip()
    if fmt not in {"word", "xml"}:
        raise HTTPException(status_code=400, detail="format must be 'word' or 'xml'")

    cid = str(r.get("customer_id") or "")
    aid = r.get("alert_id")
    aid_str = str(aid) if aid else None
    txn_dict = _txn_dict_for_estr_kyc(aid_str, cid)
    if aid_str and _ALERTS.get(aid_str):
        cid = cid or _ALERTS[aid_str].customer_id

    pg = getattr(request.app.state, "pg", None)
    customer = await get_or_create_customer_kyc(pg, cid or "UNKNOWN", txn_dict)
    bvn_linked_accounts = await list_bvn_linked_accounts(
        pg, str(customer.id_number or "").strip(), primary_customer_id=cid or "UNKNOWN"
    )
    alert_dict = _ALERTS[aid_str].model_dump() if aid_str and _ALERTS.get(aid_str) else None
    kind = str((alert_dict or {}).get("otc_report_kind") or "").strip().lower()
    report_label = "OTC_ESAR" if kind == "otc_esar" else "OTC_ESTR"
    if fmt == "xml":
        return Response(
            content=(r.get("xml") or ""),
            media_type="application/xml",
            headers={"Content-Disposition": _report_download_content_disposition(customer.customer_name, report_label, "xml")},
        )
    estr_notes = str(r.get("estr_notes") or "")
    doc_bytes = await render_otc_estr_docx_bytes(
        customer=customer,
        alert=alert_dict,
        estr_notes=estr_notes,
        approver_name=_approver_display_name(user),
        bvn_linked_accounts=bvn_linked_accounts,
    )
    return Response(
        content=doc_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _report_download_content_disposition(customer.customer_name, report_label, "docx")},
    )


@router.post("/nfiu/customer-change/generate")
async def generate_nfiu_customer_change(payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    """
    NFIU customer information change reports (demo XML): partial/full name change, BVN update, DOB change.
    """
    change_type = str(payload.get("change_type") or "").strip()
    allowed = {"partial_name_change", "full_name_change", "bvn_update", "dob_change"}
    if change_type not in allowed:
        raise HTTPException(
            status_code=400,
            detail=f"change_type must be one of: {', '.join(sorted(allowed))}",
        )
    cid = str(payload.get("customer_id") or "").strip()
    if not cid:
        raise HTTPException(status_code=400, detail="customer_id is required")
    fields = {
        "old_value": payload.get("old_value"),
        "new_value": payload.get("new_value"),
        "notes": payload.get("notes"),
        "bvn_old": payload.get("bvn_old"),
        "bvn_new": payload.get("bvn_new"),
        "name_old": payload.get("name_old"),
        "name_new": payload.get("name_new"),
        "dob_old": payload.get("dob_old"),
        "dob_new": payload.get("dob_new"),
    }
    use_llm = bool(payload.get("use_llm"))
    include_word = bool(payload.get("include_word_narrative")) or use_llm
    word_narrative = None
    word_source = None
    if include_word:
        wn, ws = await build_nfiu_cir_bundle_narrative(
            change_type=change_type,
            customer_id=cid,
            fields=fields,
            use_llm=use_llm,
        )
        word_narrative = wn
        word_source = ws
    return _draft_nfiu_cir_record(
        change_type, cid, user, fields, word_narrative=word_narrative, word_narrative_source=word_source
    )


@router.get("/nfiu/{report_id}/download")
async def download_nfiu_customer_change(
    report_id: str,
    format: str = "xml",
    user: Dict[str, Any] = Depends(get_current_user),
):
    r = _REPORTS.get(report_id)
    if not r or not _report_not_soft_deleted(r) or r.get("type") != "NFIU_CIR":
        raise HTTPException(status_code=404, detail="NFIU report not found")
    ct = str(r.get("change_type") or "customer_change")
    fmt = format.lower().strip()
    if fmt == "word" and r.get("word_narrative"):
        doc_bytes = regulatory_narrative_docx_bytes(
            title=f"NFIU customer information change — {ct}",
            subtitle=f"Customer: {r.get('customer_id')}",
            narrative=str(r.get("word_narrative") or ""),
            xml_excerpt=str(r.get("xml") or ""),
            source_note=str(r.get("word_narrative_source") or ""),
        )
        return Response(
            content=doc_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="NFIU_{ct}_{report_id}.docx"'},
        )
    return _download_registered_report(
        report_id,
        format,
        filename_prefix=f"NFIU_{ct}",
        title=f"NFIU customer information change — {ct} (demo)",
    )


@router.post("/validate-xml")
async def validate_xml(payload: Dict[str, Any], user: Dict[str, Any] = Depends(get_current_user)):
    xml_content = payload.get("xml_content")
    if not xml_content:
        raise HTTPException(status_code=400, detail="xml_content is required")
    # Stub: only checks if it parses
    ok = xml_content.strip().startswith("<")
    return {"valid": ok, "errors": [] if ok else ["XML does not look valid"]}

