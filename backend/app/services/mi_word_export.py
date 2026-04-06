"""Board pack Word summary (python-docx); aligns with templates/reports/board_pack_title.txt."""

from __future__ import annotations

import io
from typing import Any, Dict

from docx import Document


def build_board_pack_docx_bytes(payload: Dict[str, Any]) -> bytes:
    doc = Document()
    title = payload.get("template_title") or "Board AML Management Information Pack"
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generated (UTC): {payload.get('generated_at', '')}")
    kpi = payload.get("kpi") or {}
    doc.add_heading("Key indicators", level=1)
    for label, key in [
        ("Total alerts", "total_alerts"),
        ("STR submitted", "str_filed_submitted"),
        ("Open >7 days", "open_alerts_ageing_over_7_days"),
        ("Material escalations", "material_escalations_count"),
        ("Pending CCO STR approvals", "pending_cco_str_approvals"),
    ]:
        doc.add_paragraph(f"{label}: {kpi.get(key, '—')}", style="List Bullet")
    doc.add_paragraph(payload.get("disclaimer") or "")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
