"""
Suspicious Activity Report (SAR) — Microsoft Word layout aligned to internal bank template
(sections I–VII: Profile through Approval).
"""
from __future__ import annotations

import json
import os
import re
from datetime import date, datetime
from io import BytesIO
from typing import Any, Dict, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.core.logging import get_logger
from app.services.str_word_generator import CustomerKyc, _date_to_long, _format_money_with_currency

log = get_logger(component="sar_word_generator")

# JSON keys produced by the LLM / template fallback — matches Word template subsection content.
_SAR_SCHEMA_KEYS: Tuple[str, ...] = (
    "primary_suspicion",
    "secondary_suspicion",
    "activity_typology",
    "bvn_internal_linkage",
    "bvn_other_accounts",
    "bvn_volume_comparison",
    "chronology_text",
    "nfiu_finding_1_title",
    "nfiu_finding_1_body",
    "nfiu_finding_2_title",
    "nfiu_finding_2_body",
    "nfiu_finding_3_title",
    "nfiu_finding_3_body",
    "action_internal_review",
    "action_monitoring",
    "action_reporting",
)

_LEGACY_SECTION_KEYS = frozenset(
    {
        "activity_classification",
        "bvn_relationship_analytics",
        "chronology",
        "nfiu_narrative_investigation",
        "action_taken",
    }
)

_LLM_SYSTEM = """You are a senior Nigerian bank AML reporting officer drafting a Suspicious Activity Report (SAR)
for internal and NFIU-facing use. The case may be activity-led (false-positive closure / pattern SAR) or tied to
OTC ESAR identity matters (see activity_profile.workflow == "otc_esar" and otc_* fields when present).

If us_activity_focus is true in the context, you may note US-person, US financial institution, or USD-clearing nexus
only when plausible from the data; otherwise do not emphasise US nexus.

Output a single JSON object only (no markdown fences). Keys must be exactly these strings, each value a non-empty string:
primary_suspicion, secondary_suspicion, activity_typology,
bvn_internal_linkage, bvn_other_accounts, bvn_volume_comparison,
chronology_text,
nfiu_finding_1_title, nfiu_finding_1_body, nfiu_finding_2_title, nfiu_finding_2_body, nfiu_finding_3_title, nfiu_finding_3_body,
action_internal_review, action_monitoring, action_reporting.

Rules for CONTENT (not the JSON keys):
- Do NOT repeat the field labels in the values (e.g. primary_suspicion value should be the suspicion text only, not "Primary Suspicion: ...").
- activity_typology: one concise typology line (template style).
- chronology_text: several short paragraphs/lines matching a professional SAR. Use structure like:
  "Inflow Phase: ..." then optional sub-lines (e.g. origin/beneficiary hints from context), then "Outflow Phase: ..." when relevant.
  Use ₦ for naira. Reference dates from transaction.date_display / created_at when present.
- nfiu_finding_*_title: short heading text only WITHOUT a leading number (the document adds "1.", "2.", "3.").
  Choose titles that fit the scenario (e.g. "Behavioral Deviation", "Source of Funds Concerns", "Relationship Substantiation"
  for pass-through patterns; adapt for OTC ESAR / identity matters).
- nfiu_finding_*_body: 1–3 sentences each, formal English.
- action_*: single-sentence descriptions (content after "Internal Review:", "Monitoring:", "Reporting:" labels in the final doc).

Tailor every section to the scenario in the JSON context; avoid generic filler when the context gives customer, amounts, narratives, or OTC details."""


def _txn_date_long(txn: Dict[str, Any]) -> str:
    ts = txn.get("created_at") or txn.get("timestamp")
    if isinstance(ts, datetime):
        return ts.strftime("%B %d, %Y")
    if isinstance(ts, str):
        try:
            return datetime.fromisoformat(ts.replace("Z", "+00:00")).strftime("%B %d, %Y")
        except Exception:
            pass
    return datetime.utcnow().strftime("%B %d, %Y")


def _clean_llm_json(text: str) -> str:
    t = text.strip()
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z0-9]*\s*", "", t)
        t = re.sub(r"\s*```\s*$", "", t)
    return t.strip()


def _parse_llm_sections(raw: str) -> Optional[Dict[str, str]]:
    try:
        data = json.loads(_clean_llm_json(raw))
    except json.JSONDecodeError:
        return None
    if not isinstance(data, dict):
        return None
    out: Dict[str, str] = {}
    for k in _SAR_SCHEMA_KEYS:
        v = data.get(k)
        if not isinstance(v, str) or not v.strip():
            return None
        out[k] = v.strip()
    return out


def _us_nexus_clause(us_focus: bool) -> str:
    if not us_focus:
        return ""
    return (
        " Cross-border exposure: assess any plausible US-person, US institution, or USD-clearing nexus from available metadata (demo)."
    )


def _bvn_other_accounts_text(ctx: Dict[str, Any]) -> str:
    cust = ctx.get("customer") if isinstance(ctx.get("customer"), dict) else {}
    primary_acct = str(cust.get("account_number") or "").strip()
    linked = ctx.get("bvn_linked_accounts") if isinstance(ctx.get("bvn_linked_accounts"), list) else []
    accounts: list[str] = []
    for row in linked:
        if not isinstance(row, dict):
            continue
        acct = str(row.get("account_number") or "").strip()
        if acct:
            accounts.append(acct)
    uniq = sorted(set(accounts))
    other = [a for a in uniq if a != primary_acct]
    if not other:
        return "No other accounts linked to this BVN were identified in the Bank's records."
    return ", ".join(other)


def _fallback_sections(ctx: Dict[str, Any]) -> Dict[str, str]:
    cust = ctx.get("customer") or {}
    txn = ctx.get("transaction") or {}
    alert = ctx.get("alert") or {}
    ap = ctx.get("activity_profile") if isinstance(ctx.get("activity_profile"), dict) else {}
    us_focus = bool(ctx.get("us_activity_focus"))
    usx = _us_nexus_clause(us_focus)

    name = cust.get("customer_name") or "the customer"
    acct = cust.get("account_number") or "—"
    bvn = cust.get("id_number") or cust.get("bvn") or "—"
    occ = cust.get("line_of_business") or "—"
    amt = float(txn.get("amount") or 0)
    amt_s = _format_money_with_currency(amt) if amt else "₦0"
    narr = str(txn.get("narrative") or "unspecified narrative")[:400]
    summ = str(alert.get("summary") or ap.get("alert_summary") or "suspicious activity indicators")[:320]
    tdate = txn.get("date_display") or _txn_date_long(txn)
    fp_notes = str(ap.get("false_positive_resolution_notes") or "")[:600]
    hints = ap.get("scenario_hints") or []
    hints_txt = "; ".join(str(x) for x in hints if x)[:500]

    wf = str(ap.get("workflow") or "").strip()

    if wf == "otc_esar":
        subj = str(ap.get("otc_subject") or "OTC identity matter")[:280]
        rationale = str(ap.get("otc_officer_rationale") or ap.get("otc_filing_reason_detail") or "")[:600]
        return {
            "primary_suspicion": f"Identity / profile inconsistency and OTC ESAR matter: {subj}.{usx}".strip(),
            "secondary_suspicion": "Deviation from verified KYC attributes and insufficient substantiation of stated occupation or source of wealth relative to observed risk flags.",
            "activity_typology": f"OTC ESAR — identity and profile alignment risk (linked to alert typology: {summ[:200]}).",
            "bvn_internal_linkage": (
                f"The subject's BVN ({bvn}) maps to account {acct} within the institution. Occupation on file: {occ}."
            ),
            "bvn_other_accounts": _bvn_other_accounts_text(ctx),
            "bvn_volume_comparison": (
                f"Recent referenced activity includes {amt_s} with narrative context from monitoring; compare to historical low-value profile where applicable."
            ),
            "chronology_text": (
                f"Detection phase: Monitoring and OTC workflow documented the matter: {subj}.\n"
                f"Review phase: Officer rationale captured in the case file: {rationale or 'See OTC filing record (demo).'}\n"
                f"Anchor event: {tdate} — transaction narrative reference: {narr}"
            ),
            "nfiu_finding_1_title": "Identity and profile alignment",
            "nfiu_finding_1_body": (
                f"The customer's stated profile and KYC data do not fully align with the risk indicators summarised in the alert: {summ[:240]}."
            ),
            "nfiu_finding_2_title": "OTC officer assessment",
            "nfiu_finding_2_body": (rationale[:800] or "Officer rationale is recorded on the OTC ESAR track; SAR documents the identity-related suspicion for regulatory awareness (demo)."),
            "nfiu_finding_3_title": "Regulatory posture",
            "nfiu_finding_3_body": "Filing documents the identity-related suspicion and investigation summary consistent with internal policy and NFIU reporting expectations (demo).",
            "action_internal_review": "Reviewed KYC refresh, OTC worksheets, and alert investigation history.",
            "action_monitoring": "Account retained under enhanced monitoring consistent with risk appetite (demo).",
            "action_reporting": "This SAR is being filed with the Nigeria Financial Intelligence Unit (NFIU) to comply with AML/CFT obligations.",
        }

    if ap:
        return {
            "primary_suspicion": (
                f"Suspicious activity pattern inconsistent with known occupation and expected behaviour: {summ[:220]}.{usx}"
            ).strip(),
            "secondary_suspicion": (
                "Deviation from historical behaviour and typology signals across the activity window; economic purpose unclear from available records."
            ),
            "activity_typology": (
                (f"Pattern-led SAR after false-positive STR disposition; context hints: {hints_txt}. " if hints_txt else "")
                + "Activity narrative emphasises scenario / typology rather than a single ledger row (demo)."
            ).strip(),
            "bvn_internal_linkage": f"BVN ({bvn}) maps to account {acct}; relationship view emphasises behavioural signals over a single posting.",
            "bvn_other_accounts": _bvn_other_accounts_text(ctx),
            "bvn_volume_comparison": (
                f"Referenced anchor amount {amt_s} (illustrative); SAR focuses on broader activity described in monitoring output."
                if amt
                else "Volume assessment relies on typology and investigation notes rather than a single transaction."
            ),
            "chronology_text": (
                "Detection phase: Rules engine surfaced the activity described in the alert summary.\n"
                f"Reference narrative: {narr}\n"
                + (
                    f"Transaction-linked anchor (if any): {tdate} — {amt_s}.\n"
                    if amt
                    else "Activity window: consolidate timelines from case files and monitoring hits.\n"
                )
                + f"Post-review notes: {fp_notes or 'Documented under false-positive closure; SAR captures residual activity context (demo).'}"
            ),
            "nfiu_finding_1_title": "Suspicious activity (pattern-led)",
            "nfiu_finding_1_body": (
                f"The observed behaviour for {name} aligns with scenario indicators: {summ[:280]}."
            ),
            "nfiu_finding_2_title": "Investigation summary",
            "nfiu_finding_2_body": fp_notes or "Investigation notes are summarised from the alert workflow and typology hits (demo).",
            "nfiu_finding_3_title": "Regulatory posture",
            "nfiu_finding_3_body": "Filing documents the activity narrative for NFIU awareness consistent with internal policy (demo).",
            "action_internal_review": "Comprehensive review of KYC, income profile, and investigation record supporting the false-positive closure track.",
            "action_monitoring": "Retain enhanced monitoring settings as per risk appetite (demo).",
            "action_reporting": "This SAR is being filed with the Nigeria Financial Intelligence Unit (NFIU) to comply with AML/CFT regulatory obligations.",
        }

    return {
        "primary_suspicion": (
            f"Inconsistency with known occupation ({occ}) and income profile relative to flagged activity.{usx}"
        ).strip(),
        "secondary_suspicion": "Deviation from historical customer behaviour and lack of clear economic purpose for the observed flow.",
        "activity_typology": f"Aligned with alert summary — {summ[:220]}.",
        "bvn_internal_linkage": f"The subject's BVN ({bvn}) is linked to account {acct} within the institution.",
        "bvn_other_accounts": _bvn_other_accounts_text(ctx),
        "bvn_volume_comparison": (
            f"Recent activity ({amt_s}) represents a material step-change relative to prior lower-value activity on file."
        ),
        "chronology_text": (
            f"Inflow phase: On {tdate}, activity on the account included the flagged movement ({amt_s}).\n"
            f"Narrative reference: {narr}\n"
            "Outflow phase: Review internal ledgers for onward transfers following the credit where applicable — see operations notes."
        ),
        "nfiu_finding_1_title": "Behavioral deviation",
        "nfiu_finding_1_body": (
            f"The pattern associated with {name} departs from the established relationship profile on file."
        ),
        "nfiu_finding_2_title": "Source of funds / economic purpose",
        "nfiu_finding_2_body": "The economic justification for the flagged value and counterparties requires further substantiation under EDD.",
        "nfiu_finding_3_title": "Relationship substantiation",
        "nfiu_finding_3_body": "EDD was unable to fully substantiate a legitimate business rationale consistent with the observed flows.",
        "action_internal_review": "Conducted a comprehensive review of KYC, income profile, and historical patterns.",
        "action_monitoring": "The account has been placed under enhanced monitoring.",
        "action_reporting": "This SAR is being filed with the Nigeria Financial Intelligence Unit (NFIU) to comply with AML/CFT regulatory obligations.",
    }


async def generate_sar_narrative_sections(case_context: Dict[str, Any]) -> tuple[Dict[str, str], str]:
    """Returns (sections_dict, source_label) where source_label is 'llm' or 'template'."""
    ctx_json = json.dumps(case_context, default=str, indent=2)[:12000]
    prompt = (
        f"Case context (JSON):\n{ctx_json}\n\n"
        "Generate the JSON string fields matching the bank SAR Word template (sections II–VI). "
        "Each section must reflect THIS scenario's customer, transaction, alert, activity_profile, and any OTC fields."
    )
    try:
        from app.services.llm.client import get_llm_client

        client = get_llm_client()
        res = await client.generate(prompt, system=_LLM_SYSTEM, temperature=0.5)
        parsed = _parse_llm_sections(res.content)
        if parsed:
            # Keep BVN-linked account numbers deterministic from snapshot context.
            parsed["bvn_other_accounts"] = _bvn_other_accounts_text(case_context)
            return parsed, "llm"
        log.warning("sar_llm_json_parse_failed")
    except Exception as exc:
        log.info("sar_llm_skipped err=%s", exc)

    return _fallback_sections(case_context), "template"


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _add_section_heading(doc: Document, roman: str, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{roman}. {title.upper()}")
    run.bold = True
    run.font.size = Pt(11)


def _add_labeled_paragraph(doc: Document, label: str, value: str) -> None:
    doc.add_paragraph(f"{label}: {value}")


def _append_formatted_block(doc: Document, body: str) -> None:
    """Plain paragraphs with bullets (legacy / chronology freeform)."""
    for line in body.split("\n"):
        raw = line.rstrip()
        if not raw.strip():
            doc.add_paragraph("")
            continue
        s = raw.strip()
        if s.startswith("o ") or s.startswith("○ "):
            doc.add_paragraph(f"   ○   {s[2:].strip()}")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(s)
        elif s.startswith("•"):
            doc.add_paragraph(f"•\t{s[1:].strip()}")
        elif s.startswith("-"):
            doc.add_paragraph(f"•\t{s[1:].strip()}")
        else:
            doc.add_paragraph(s)


def _add_nfiu_finding(doc: Document, index: int, title: str, body: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{index}. {title.strip()}")
    run.bold = True
    run.font.size = Pt(11)
    for block in body.split("\n\n"):
        b = block.strip()
        if b:
            doc.add_paragraph(b)


def _insert_signature_image(doc: Document, image_path: Optional[str], *, max_width_in: float = 2.2) -> None:
    path = (image_path or "").strip()
    if path and os.path.isfile(path):
        try:
            doc.add_picture(path, width=Inches(max_width_in))
            return
        except Exception as exc:
            log.warning("sar_signature_image_failed path=%s err=%s", path, exc)
    p = doc.add_paragraph("_______________________________")
    p.paragraph_format.space_after = Pt(6)


def _sar_sections_are_legacy(sections: Dict[str, Any]) -> bool:
    if not isinstance(sections, dict):
        return False
    if "primary_suspicion" in sections:
        return False
    return bool(_LEGACY_SECTION_KEYS.intersection(sections.keys()))


def _render_sar_docx_legacy(
    doc: Document,
    *,
    customer: CustomerKyc,
    sections: Dict[str, str],
    approver_name: str,
    signature_image_path: Optional[str],
    report_date: date,
    narrative_source: str,
) -> None:
    """Pre-template SAR shape (five combined blocks)."""
    _add_title(doc, "SUSPICIOUS ACTIVITY REPORT (SAR)")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Note: This file uses a legacy section layout; regenerate the SAR for the current template."
        if narrative_source != "llm"
        else "Note: Legacy layout export."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    doc.add_paragraph("")

    _add_section_heading(doc, "I", "PROFILE")
    doc.add_paragraph(f"•\tCustomer Name: {customer.customer_name}")
    doc.add_paragraph(f"•\tAccount Number: {customer.account_number}")
    doc.add_paragraph(f"•\tBVN / ID Number: {customer.id_number}")
    doc.add_paragraph(f"•\tOccupation: {customer.line_of_business}")
    doc.add_paragraph(f"•\tAddress: {customer.customer_address}")
    doc.add_paragraph(f"•\tRelationship Start Date: {_date_to_long(customer.account_opened)}")

    _add_section_heading(doc, "II", "ACTIVITY CLASSIFICATION")
    _append_formatted_block(doc, sections.get("activity_classification") or "")

    _add_section_heading(doc, "III", "BVN & RELATIONSHIP ANALYTICS")
    _append_formatted_block(doc, sections.get("bvn_relationship_analytics") or "")

    _add_section_heading(doc, "IV", "CHRONOLOGY OF SUSPICIOUS ACTIVITY")
    _append_formatted_block(doc, sections.get("chronology") or "")

    _add_section_heading(doc, "V", "NFIU NARRATIVE & INVESTIGATION FINDINGS")
    _append_formatted_block(doc, sections.get("nfiu_narrative_investigation") or "")

    _add_section_heading(doc, "VI", "ACTION TAKEN")
    _append_formatted_block(doc, sections.get("action_taken") or "")

    _add_section_heading(doc, "VII", "APPROVAL")
    doc.add_paragraph(f"Approver: {approver_name.strip() or 'Chief Compliance Officer'}")
    sig_label = doc.add_paragraph()
    sig_label.add_run("Signature: ").bold = True
    _insert_signature_image(doc, signature_image_path)
    doc.add_paragraph(f"Date: {report_date.strftime('%B %d, %Y')}")


def build_sar_case_context(
    *,
    customer: CustomerKyc,
    txn: Dict[str, Any],
    alert: Optional[Dict[str, Any]],
    sar_notes: str,
    enrichment: Optional[Dict[str, Any]],
    activity_profile: Optional[Dict[str, Any]] = None,
    us_activity_focus: bool = False,
) -> Dict[str, Any]:
    ctx: Dict[str, Any] = {
        "customer": {
            "customer_name": customer.customer_name,
            "account_number": customer.account_number,
            "id_number": customer.id_number,
            "line_of_business": customer.line_of_business,
            "customer_address": customer.customer_address,
            "phone_number": customer.phone_number,
            "account_opened": customer.account_opened.isoformat(),
            "date_of_birth": customer.date_of_birth.isoformat(),
        },
        "transaction": {
            "id": txn.get("id"),
            "amount": txn.get("amount"),
            "currency": txn.get("currency") or "NGN",
            "transaction_type": txn.get("transaction_type"),
            "narrative": txn.get("narrative"),
            "created_at": str(txn.get("created_at") or ""),
            "date_display": _txn_date_long(txn),
            "counterparty_name": txn.get("counterparty_name"),
            "counterparty_id": txn.get("counterparty_id"),
        },
        "sar_notes": sar_notes.strip(),
        "us_activity_focus": us_activity_focus,
        "narrative_basis": "suspicious_activity" if activity_profile else "transaction_led",
    }
    if activity_profile:
        ctx["activity_profile"] = activity_profile
    if alert:
        ad: Dict[str, Any] = {
            "id": alert.get("id"),
            "summary": alert.get("summary"),
            "severity": alert.get("severity"),
            "status": alert.get("status"),
            "rule_ids": alert.get("rule_ids"),
            "last_resolution": alert.get("last_resolution"),
        }
        for k in (
            "otc_subject",
            "otc_report_kind",
            "otc_outcome",
            "otc_filing_reason",
            "otc_filing_reason_detail",
            "otc_officer_rationale",
        ):
            if alert.get(k) is not None:
                ad[k] = alert.get(k)
        ctx["alert"] = ad
    if enrichment:
        why = enrichment.get("why_suspicious") if isinstance(enrichment, dict) else None
        if why:
            ctx["typology_context"] = why
        tx = enrichment.get("transaction")
        if isinstance(tx, dict):
            ctx["snapshot_transaction"] = {
                "narrative": tx.get("narrative"),
                "amount": tx.get("amount"),
            }
        bvn_linked = enrichment.get("bvn_linked_accounts")
        if isinstance(bvn_linked, list):
            ctx["bvn_linked_accounts"] = bvn_linked
    return ctx


def render_sar_docx_bytes(
    *,
    customer: CustomerKyc,
    txn: Dict[str, Any],
    sections: Dict[str, str],
    approver_name: str,
    signature_image_path: Optional[str],
    report_date: Optional[date] = None,
    narrative_source: str = "template",
) -> bytes:
    doc = Document()
    rd = report_date or date.today()
    sec = sections or {}

    if _sar_sections_are_legacy(sec):
        _render_sar_docx_legacy(
            doc,
            customer=customer,
            sections=sec,
            approver_name=approver_name,
            signature_image_path=signature_image_path,
            report_date=rd,
            narrative_source=narrative_source,
        )
        out = BytesIO()
        doc.save(out)
        return out.getvalue()

    _add_title(doc, "SUSPICIOUS ACTIVITY REPORT")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src = "AI-assisted draft" if narrative_source == "llm" else "Template-based draft"
    nr = note.add_run(
        f"Sections II–VI follow the standard SAR layout (section I Profile and section VII Approval are system fields; {src}). "
        "Review before external use."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    doc.add_paragraph("")

    _add_section_heading(doc, "I", "PROFILE")
    _add_labeled_paragraph(doc, "Customer Name", customer.customer_name)
    _add_labeled_paragraph(doc, "Account Number", customer.account_number)
    _add_labeled_paragraph(doc, "BVN / ID Number", customer.id_number)
    _add_labeled_paragraph(doc, "Occupation", customer.line_of_business)
    _add_labeled_paragraph(doc, "Address", customer.customer_address)
    _add_labeled_paragraph(doc, "Relationship Start Date", _date_to_long(customer.account_opened))
    doc.add_paragraph("")

    _add_section_heading(doc, "II", "ACTIVITY CLASSIFICATION")
    _add_labeled_paragraph(doc, "Primary Suspicion", sec.get("primary_suspicion") or "—")
    _add_labeled_paragraph(doc, "Secondary Suspicion", sec.get("secondary_suspicion") or "—")
    _add_labeled_paragraph(doc, "Typology", sec.get("activity_typology") or "—")
    doc.add_paragraph("")

    _add_section_heading(doc, "III", "BVN & RELATIONSHIP ANALYTICS")
    _add_labeled_paragraph(doc, "Internal Linkage", sec.get("bvn_internal_linkage") or "—")
    _add_labeled_paragraph(doc, "Other Accounts", sec.get("bvn_other_accounts") or "—")
    _add_labeled_paragraph(doc, "Volume Comparison", sec.get("bvn_volume_comparison") or "—")
    doc.add_paragraph("")

    _add_section_heading(doc, "IV", "CHRONOLOGY OF SUSPICIOUS ACTIVITY (DEMO)")
    _append_formatted_block(doc, sec.get("chronology_text") or "")
    doc.add_paragraph("")

    _add_section_heading(doc, "V", "NFIU NARRATIVE & INVESTIGATION FINDINGS")
    _add_nfiu_finding(doc, 1, sec.get("nfiu_finding_1_title") or "Finding 1", sec.get("nfiu_finding_1_body") or "—")
    doc.add_paragraph("")
    _add_nfiu_finding(doc, 2, sec.get("nfiu_finding_2_title") or "Finding 2", sec.get("nfiu_finding_2_body") or "—")
    doc.add_paragraph("")
    _add_nfiu_finding(doc, 3, sec.get("nfiu_finding_3_title") or "Finding 3", sec.get("nfiu_finding_3_body") or "—")
    doc.add_paragraph("")

    _add_section_heading(doc, "VI", "ACTION TAKEN")
    _add_labeled_paragraph(doc, "Internal Review", sec.get("action_internal_review") or "—")
    _add_labeled_paragraph(doc, "Monitoring", sec.get("action_monitoring") or "—")
    _add_labeled_paragraph(doc, "Reporting", sec.get("action_reporting") or "—")
    doc.add_paragraph("")

    _add_section_heading(doc, "VII", "APPROVAL")
    doc.add_paragraph(f"Approver: {approver_name.strip() or 'Chief Compliance Officer'}")
    sig_label = doc.add_paragraph()
    sig_label.add_run("Signature: ").bold = True
    _insert_signature_image(doc, signature_image_path)
    doc.add_paragraph(f"Date: {rd.strftime('%B %d, %Y')}")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
