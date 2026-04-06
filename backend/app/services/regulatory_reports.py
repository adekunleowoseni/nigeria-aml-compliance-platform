"""Stub goAML / NFIU-style report payloads for demo (not production filing)."""
from __future__ import annotations

import xml.etree.ElementTree as ET
from datetime import datetime
from io import BytesIO
from typing import Any, Dict
from uuid import uuid4
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer


def _el(parent: ET.Element, tag: str, text: str | None = None) -> ET.Element:
    e = ET.SubElement(parent, tag)
    if text is not None:
        e.text = str(text)
    return e


def goaml_stub_xml(report_kind: str, payload: Dict[str, Any]) -> str:
    root = ET.Element("goAML", {"xmlns": "demo", "version": "1.0", "kind": report_kind})
    _el(root, "ReportKind", report_kind)
    _el(root, "ReportId", payload.get("report_id") or str(uuid4()))
    _el(root, "GeneratedAt", datetime.utcnow().isoformat() + "Z")
    ent = ET.SubElement(root, "ReportingEntity")
    _el(ent, "Name", payload.get("entity_name", "Demo Reporting Entity"))
    _el(ent, "Registration", payload.get("entity_rc", "RC-000000"))
    for k, v in payload.items():
        if k in ("report_id", "entity_name", "entity_rc"):
            continue
        _el(root, k, str(v)[:2000])
    return ET.tostring(root, encoding="unicode", method="xml")


def nfiu_customer_change_xml(change_type: str, fields: Dict[str, Any]) -> str:
    root = ET.Element("NFIUCustomerInformationChange", {"version": "demo"})
    _el(root, "ChangeType", change_type)
    _el(root, "ReportReference", fields.get("report_id") or str(uuid4()))
    _el(root, "SubmittedAt", datetime.utcnow().isoformat() + "Z")
    _el(root, "CustomerId", str(fields.get("customer_id", "")))
    for key in ("old_value", "new_value", "notes", "bvn_old", "bvn_new", "name_old", "name_new", "dob_old", "dob_new"):
        if key in fields and fields[key]:
            _el(root, key.replace("_", ""), str(fields[key])[:4000])
    return ET.tostring(root, encoding="unicode", method="xml")


def minimal_docx_bytes(title: str, body: str) -> bytes:
    d = Document()
    d.add_heading(title, 0)
    text = body or ""
    for chunk in [text[i : i + 8000] for i in range(0, min(len(text), 24000), 8000)] or [""]:
        d.add_paragraph(chunk)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def regulatory_narrative_docx_bytes(
    *,
    title: str,
    subtitle: str,
    narrative: str,
    xml_excerpt: str,
    source_note: str | None = None,
) -> bytes:
    """Structured Word doc for bundle-generated AOP / NFIU CIR (narrative + XML excerpt)."""
    d = Document()
    d.add_heading(title, 0)
    if subtitle:
        d.add_paragraph(subtitle)
    if source_note:
        p = d.add_paragraph()
        r = p.add_run(f"Narrative source: {source_note}")
        r.italic = True
    d.add_heading("Internal narrative", level=1)
    for block in (narrative or "").split("\n\n"):
        b = block.strip()
        if b:
            d.add_paragraph(b)
    d.add_heading("XML payload (excerpt)", level=1)
    xml_text = (xml_excerpt or "")[:14000]
    for chunk in [xml_text[i : i + 8000] for i in range(0, len(xml_text), 8000)] or [""]:
        d.add_paragraph(chunk)
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


def regulatory_narrative_pdf_bytes(
    *,
    title: str,
    subtitle: str,
    narrative: str,
    source_note: str | None = None,
    xml_excerpt: str | None = None,
) -> bytes:
    """Customer-facing PDF (ReportLab). Optional XML appendix for internal packs only."""
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    story: list[Any] = []
    story.append(Paragraph(f"<b>{escape(title)}</b>", styles["Title"]))
    story.append(Spacer(1, 0.3 * cm))
    if subtitle:
        story.append(Paragraph(escape(subtitle), styles["Normal"]))
        story.append(Spacer(1, 0.2 * cm))
    if source_note:
        story.append(Paragraph(f"<i>Narrative source: {escape(source_note)}</i>", styles["Normal"]))
        story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("<b>Narrative</b>", styles["Heading2"]))
    for block in (narrative or "").split("\n\n"):
        b = block.strip()
        if b:
            story.append(Paragraph(escape(b).replace("\n", "<br/>"), styles["Normal"]))
            story.append(Spacer(1, 0.15 * cm))
    if xml_excerpt:
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("<b>Technical appendix (XML excerpt)</b>", styles["Heading2"]))
        ex = (xml_excerpt or "")[:12000]
        code_style = ParagraphStyle("XmlExcerpt", parent=styles["Code"], fontName="Courier", fontSize=7, leading=9)
        story.append(Preformatted(ex, code_style, maxLineLength=100))
    doc.build(story)
    return buf.getvalue()
