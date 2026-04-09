"""
OTC Extended STR (ESTR) — Microsoft Word layout aligned to internal ESTR template.

Document title branches on OTC subject (compliance officer selection):
- Cash deposit / withdrawal → OTC SUSPICIOUS TRANSACTION REPORT
- Identity / profile change matters → OTC SUSPICIOUS ACTIVITY REPORT
"""
from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.logging import get_logger
from app.models.alert import OTC_SUBJECTS_ESTR
from app.services.str_word_generator import CustomerKyc, _date_to_long

log = get_logger(component="estr_word_generator")

# Strip internal spreadsheet reference from narratives so it does not appear in regulatory Word output.
_REF_STR_SPREADSHEET = re.compile(
    r"\s*\|\s*Reference STR ID \(spreadsheet\):\s*\d+\s*",
    re.IGNORECASE,
)


def _sanitize_text_for_estr_word(s: str) -> str:
    if not (s or "").strip():
        return ""
    return _REF_STR_SPREADSHEET.sub("", s).strip()

_OTC_SUBJECT_LABELS: Dict[str, str] = {
    "cash_deposit": "Cash deposit",
    "cash_withdrawal": "Cash withdrawal",
    "change_of_name": "Change of name",
    "arrangement_of_name": "Arrangement of name",
    "nin_update": "NIN update / change",
    "bvn_partial_name_change": "BVN partial name change",
    "full_name_change": "Full name change",
    "dob_update": "Date of birth update",
    "name_and_dob_update": "Name and date of birth update",
}

_FILING_REASON_LABELS: Dict[str, str] = {
    "regulatory_obligation": "Regulatory obligation",
    "internal_policy": "Internal policy",
    "branch_referral": "Branch referral",
    "customer_request": "Customer request",
    "supervisory_request": "Supervisory / regulatory request",
    "other": "Other (specified in detail)",
}

_LLM_SYSTEM = """You are a senior Nigerian bank AML compliance officer drafting Section II ("Reasons for Filing") of an
over-the-counter (OTC) extended regulatory return (ESTR-style) for the Nigeria Financial Intelligence Unit (NFIU) / goAML context.

Requirements:
- The "Nature of Unusual Activity" in Section I is already fixed from the OTC subject and branch intake. Your job is to write
  a single polished Reasons for Filing narrative that is grounded primarily in that nature line and the OTC matter—not a
  list of form fields.
- Rewrite and clarify; do not concatenate labels (e.g. do not write "Officer selected X. Officer selected Y."). Integrate
  filing basis, rationale, and any analyst extension notes into fluent prose.
- If extension notes are provided, fold them in naturally. Do not add a separate "Additional notes" paragraph that only
  repeats the same facts unless genuinely needed for clarity.
- Align with CBN AML/CFT supervisory expectations: factual, cautious, professional; describe observations and internal
  control steps; do not assert criminal guilt.
- Output plain English only: 2–4 short paragraphs, no markdown, no numbered lists unless essential.
- Do not invent court outcomes or regulatory decisions not implied by the inputs."""


def otc_subject_display(subject: Optional[str]) -> str:
    s = (subject or "").strip().lower()
    return _OTC_SUBJECT_LABELS.get(s, (subject or "Not specified").replace("_", " ").title())


def otc_filing_reason_display(reason: Optional[str]) -> str:
    r = (reason or "").strip().lower()
    return _FILING_REASON_LABELS.get(r, (reason or "Not specified").replace("_", " ").title())


def nature_of_unusual_activity_from_otc(alert: Optional[Dict[str, Any]]) -> str:
    """
    Section I line aligned to OTC intake fields only:
    subject + filing reason (+ optional reason detail).
    Avoids typology/red-flag summary language.
    """
    if not alert:
        return "Not specified"
    subj_raw = alert.get("otc_subject")
    subject_label = otc_subject_display(str(subj_raw) if subj_raw else None)
    reason_label = otc_filing_reason_display(alert.get("otc_filing_reason"))
    detail = str(alert.get("otc_filing_reason_detail") or "").strip()
    if detail:
        if len(detail) > 180:
            detail = detail[:177].rstrip() + "…"
        return f"{subject_label} ({reason_label}) — {detail}"
    return f"{subject_label} ({reason_label})"


def otc_estr_document_title(otc_subject: Optional[str]) -> str:
    s = (otc_subject or "").strip().lower()
    if not s or s in OTC_SUBJECTS_ESTR:
        return "OTC SUSPICIOUS TRANSACTION REPORT"
    return "SUSPICIOUS ACTIVITY REPORT"


def _reporting_line(is_transaction_report: bool) -> str:
    if is_transaction_report:
        return (
            "This STR is being filed with the Nigeria Financial Intelligence Unit (NFIU) to comply with AML/CFT "
            "regulatory obligations."
        )
    return (
        "This SAR is being filed with the Nigeria Financial Intelligence Unit (NFIU) to comply with AML/CFT "
        "regulatory obligations."
    )


def _fallback_reasons_body(
    *,
    nature_label: str,
    reason_label: str,
    reason_detail: str,
    officer_rationale: str,
    estr_notes: str,
    customer_name: str,
) -> str:
    detail = (reason_detail or "").strip()
    rat = (officer_rationale or "").strip()
    ext = (estr_notes or "").strip()
    parts = [
        f"The institution is filing this OTC return following identification of unusual activity categorised as: {nature_label}. "
        f"The compliance officer selected the following basis for filing: {reason_label}."
    ]
    if detail:
        parts.append(f"Additional context recorded for the filing basis: {detail}")
    if rat:
        parts.append(f"Officer assessment and rationale: {rat}")
    if ext:
        parts.append(f"Supplementary extension notes: {ext}")
    parts.append(
        f"The subject customer ({customer_name}) is known to the bank; the narrative above reflects the current "
        "understanding pending completion of enhanced due diligence and any required regulatory follow-up."
    )
    return "\n\n".join(parts)


async def refine_estr_reasons_for_filing(
    *,
    nature_line_section_i: str,
    reason_label: str,
    reason_detail: str,
    officer_rationale: str,
    estr_notes: str,
    customer_name: str,
) -> Tuple[str, str]:
    """Returns (body_text, source) where source is 'llm' or 'template'."""
    prompt = (
        f"Customer name: {customer_name}\n"
        f"Nature of Unusual Activity (exact Section I line): {nature_line_section_i}\n"
        f"Declared filing basis (category): {reason_label}\n"
        f"Filing basis detail (if any): {reason_detail or 'None'}\n"
        f"Officer rationale: {officer_rationale or 'None'}\n"
        f"Optional analyst refinement notes (integrate, do not tack on as a separate appendix): {estr_notes or 'None'}\n\n"
        "Write Section II — Reasons for Filing only, as continuous professional prose."
    )
    try:
        from app.config import settings
        from app.services.llm.client import GeminiClient, get_llm_client

        if settings.gemini_api_key:
            client = GeminiClient(api_key=settings.gemini_api_key, model=settings.gemini_model)
        else:
            client = get_llm_client()
        res = await client.generate(prompt=prompt, system=_LLM_SYSTEM, temperature=0.35)
        body = (res.content or "").strip()
        body = re.sub(r"```[a-zA-Z]*\s*", "", body)
        body = body.strip()
        if len(body) > 80:
            return body, "llm"
    except Exception as exc:
        log.info("estr_reasons_llm_skipped err=%s", exc)

    return (
        _fallback_reasons_body(
            nature_label=nature_line_section_i,
            reason_label=reason_label,
            reason_detail=reason_detail,
            officer_rationale=officer_rationale,
            estr_notes=estr_notes,
            customer_name=customer_name,
        ),
        "template",
    )


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def _section_heading(doc: Document, roman: str, title: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{roman}. {title.upper()}")
    r.bold = True
    r.font.size = Pt(11)


def _labeled(doc: Document, label: str, value: str) -> None:
    doc.add_paragraph(f"{label}: {value}")


async def render_otc_estr_docx_bytes(
    *,
    customer: CustomerKyc,
    alert: Optional[Dict[str, Any]],
    estr_notes: str,
    approver_name: str,
    bvn_linked_accounts: Optional[list[dict[str, Any]]] = None,
) -> bytes:
    subj = (alert or {}).get("otc_subject") if alert else None
    subj_s = str(subj).strip().lower() if subj else ""
    title = otc_estr_document_title(str(subj) if subj else None)
    is_trx = (not subj_s) or (subj_s in OTC_SUBJECTS_ESTR)

    nature_line = nature_of_unusual_activity_from_otc(alert)
    reason_label = otc_filing_reason_display((alert or {}).get("otc_filing_reason") if alert else None)
    reason_detail = _sanitize_text_for_estr_word(str((alert or {}).get("otc_filing_reason_detail") or "").strip())
    officer_rationale = _sanitize_text_for_estr_word(str((alert or {}).get("otc_officer_rationale") or "").strip())

    reasons_body, _src = await refine_estr_reasons_for_filing(
        nature_line_section_i=nature_line,
        reason_label=reason_label,
        reason_detail=reason_detail,
        officer_rationale=officer_rationale,
        estr_notes=estr_notes.strip(),
        customer_name=customer.customer_name,
    )

    doc = Document()
    _add_title(doc, title)

    if not is_trx:
        # ESAR format mirrors the provided SAR-style stationery: simple profile lines,
        # explicit "REASONS FOR FILING SAR", "ACTION TAKEN", then approval attestation.
        _labeled(doc, "Customer Name", customer.customer_name)
        _labeled(doc, "Customer Account Number", customer.account_number)
        _labeled(doc, "Account Opened", _date_to_long(customer.account_opened))
        _labeled(doc, "Customer Address", customer.customer_address)
        _labeled(doc, "Line of Business", customer.line_of_business)
        _labeled(doc, "ID Card Number", customer.id_number)
        _labeled(doc, "Phone Number", customer.phone_number or "—")
        _labeled(doc, "Date of Birth", _date_to_long(customer.date_of_birth))
        _labeled(doc, "Nature of Unusual Activity", nature_line)

        p = doc.add_paragraph()
        p.add_run("REASONS FOR FILING SAR:").bold = True
        for block in reasons_body.split("\n\n"):
            b = block.strip()
            if b:
                doc.add_paragraph(b)

        linked_accounts: list[str] = []
        for row in bvn_linked_accounts or []:
            if not isinstance(row, dict):
                continue
            acct = str(row.get("account_number") or "").strip()
            if acct:
                linked_accounts.append(acct)
        linked_accounts = sorted(set(linked_accounts))
        other_accounts = [a for a in linked_accounts if a != customer.account_number]
        if other_accounts:
            bvn_link_line = f"The customer's BVN is linked to other account number(s) in the bank: {', '.join(other_accounts)}."
        else:
            bvn_link_line = (
                "A review of the bank's database reveals that the customer's BVN is not linked "
                "to any other account in the bank."
            )
        doc.add_paragraph(
            f"The customer commenced banking relationship with the bank on {_date_to_long(customer.account_opened)}. "
            f"The account was opened with account number {customer.account_number} and BVN {customer.id_number}. "
            f"{bvn_link_line}"
        )
        doc.add_paragraph(
            f"CDD and KYC were carried out at account opening where the customer was profiled as "
            f'"{customer.line_of_business}". There is currently no adverse media report treated as a final determination.'
        )
        doc.add_paragraph(
            "We have concerns over this account based on the observed profile-change activity and supporting rationale "
            "provided in this filing."
        )

        p = doc.add_paragraph()
        p.add_run("ACTION TAKEN: ").bold = True
        p.add_run(
            "The account is being closely monitored and has been reported to the NFIU in line with AML/CFT directives."
        )

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("APPROVAL").bold = True
        doc.add_paragraph(
            "I have reviewed and confirmed that my comments have been incorporated. "
            "I am approving the filing of an STR/SAR with the NFIU."
        )
        doc.add_paragraph("APPROVER: _______________________________")
        doc.add_paragraph(approver_name.strip() or "Chief Compliance Officer")
        doc.add_paragraph(f"DATE: {datetime.utcnow().strftime('%B %d, %Y')}")

        out = BytesIO()
        doc.save(out)
        return out.getvalue()

    _section_heading(doc, "I", "PROFILE")
    _labeled(doc, "Customer Name", customer.customer_name)
    _labeled(doc, "Account Number", customer.account_number)
    _labeled(doc, "BVN / ID Number", customer.id_number)
    _labeled(doc, "Occupation", customer.line_of_business)
    _labeled(doc, "Address", customer.customer_address)
    _labeled(doc, "Relationship Start Date", _date_to_long(customer.account_opened))
    _labeled(doc, "Date of Birth", _date_to_long(customer.date_of_birth))
    _labeled(doc, "Phone Number", customer.phone_number or "—")
    _labeled(doc, "Nature of Unusual Activity", nature_line)
    doc.add_paragraph("")

    _section_heading(doc, "II", "REASONS FOR FILING")
    for block in reasons_body.split("\n\n"):
        b = block.strip()
        if b:
            doc.add_paragraph(b)

    linked_accounts: list[str] = []
    for row in bvn_linked_accounts or []:
        if not isinstance(row, dict):
            continue
        acct = str(row.get("account_number") or "").strip()
        if acct:
            linked_accounts.append(acct)
    linked_accounts = sorted(set(linked_accounts))
    other_accounts = [a for a in linked_accounts if a != customer.account_number]
    other_accounts_line = (
        "No other accounts linked to this BVN were identified in the bank's records."
        if not other_accounts
        else ", ".join(other_accounts)
    )

    doc.add_paragraph(
        f"The customer commenced banking relationship with the bank on {_date_to_long(customer.account_opened)}. "
        f"The account is held under account number {customer.account_number} and BVN {customer.id_number}. "
        "A review of the bank's records confirms linkage details on file."
    )
    doc.add_paragraph(f"Other BVN-linked account numbers on file: {other_accounts_line}")
    doc.add_paragraph(
        f"CDD and KYC were carried out at account opening; the customer is profiled as \"{customer.line_of_business}\". "
        "There is currently no automated adverse-media hit treated as a final determination; manual validation remains required."
    )
    doc.add_paragraph("")

    _section_heading(doc, "III", "ACTION TAKEN")
    _labeled(
        doc,
        "Internal Review",
        "Enhanced review of KYC, transaction context, and OTC worksheets was completed in line with internal AML policy.",
    )
    _labeled(
        doc,
        "Monitoring",
        "The relationship is subject to enhanced monitoring and ongoing transaction surveillance as appropriate to the risk rating.",
    )
    _labeled(doc, "Reporting", _reporting_line(is_trx))
    doc.add_paragraph("")

    _section_heading(doc, "IV", "APPROVAL")
    _labeled(doc, "Approver", approver_name.strip() or "Chief Compliance Officer")
    sig = doc.add_paragraph()
    sig.add_run("Signature: ").bold = True
    sig.add_run("_______________________________")
    doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%B %d, %Y')}")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
