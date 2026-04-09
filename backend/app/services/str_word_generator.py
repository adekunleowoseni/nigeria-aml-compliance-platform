from __future__ import annotations

import hashlib
import os
import re
from dataclasses import dataclass
from datetime import datetime, date, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document

# STR Word output uses only ``New STR SUSPICIOUS TRANSACTION REPORT.docx`` stationery (repo root,
# ``backend/demo_assets``, or ``AML_STR_WORD_TEMPLATE``).
_STR_TEMPLATE_CANDIDATES = (
    "New STR SUSPICIOUS TRANSACTION REPORT.docx",
    "STR_Suspicious_Transaction_Report.docx",
    "str_word_template.docx",
)

# Compact STR body (target ~3 pages): cap list length and narrative repetition while keeping section order.
_STR_MAX_NATURE_TITLES = 4
_STR_TITLE_SOFT_MAX = 48
_STR_CLARIFY_MAX_POINTS = 2
_STR_CLARIFY_EACH_MAX = 120
_STR_TXN_FLOW_COUNTERPARTIES_MAX = 1
_STR_EVIDENCE_TITLE_ONLY_MAX = 5
_STR_CONCERN_BODY_MAX = 80
_STR_PARA_SOFT_MAX = 720
_STR_FIELD_NATURE_MAX = 420
_STR_FIELD_TXN_DESC_MAX = 900
_STR_FIELD_CLARIFY_MAX = 320


def resolve_str_word_template_path() -> Optional[Path]:
    """Resolve optional STR Word template; callers may open it to preserve typography from the bank file."""
    env = (os.environ.get("AML_STR_WORD_TEMPLATE") or "").strip()
    if env:
        p = Path(env)
        return p if p.is_file() else None
    here = Path(__file__).resolve()
    roots = (
        here.parents[2] / "demo_assets",
        here.parents[3] / "demo_assets",
        here.parents[3],
    )
    for root in roots:
        for name in _STR_TEMPLATE_CANDIDATES:
            cand = root / name
            if cand.is_file():
                return cand
    return None


# Typology / rule titles for STR "Nature of transaction" when enrichment typologies are absent.
_RULE_ID_TO_TITLE: Dict[str, str] = {
    "TYP-HUGE-INFLOW-THRESHOLD": "Large inbound credit vs policy threshold",
    "TYP-NEAR-POLICY-CEILING": "Inbound clustered below policy threshold",
    "TYP-YTD-EXCEEDS-DECLARED-TURNOVER": "YTD inflows exceed declared turnover",
    "TYP-FIRST-HUGE": "Step-change amount vs history",
    "TYP-FIRST-LARGE-INFLOW": "First large inbound credit",
    "TYP-PATTERN-INCONSISTENT": "Inbound pattern inconsistent",
    "TYP-SUDDEN-MOVEMENT": "Sudden movement / velocity cluster",
    "TYP-RAPID-INFLOW-OUTFLOW": "Rapid inbound then outbound movement",
    "TYP-DORMANT-REACTIVATION": "Dormant account reactivation",
    "TYP-FAN-IN": "Multiple inbound sources",
    "TYP-FAN-OUT": "Multiple outbound destinations",
    "TYP-STRUCTURING": "Structured/split inflows",
    "TYP-CORP-TO-INDIVIDUAL": "Corporate-to-individual flow",
    "TYP-GOV-FLOW": "Government-themed flow references",
    "TYP-PROFILE-MISMATCH": "Occupation/profile mismatch",
    "TYP-EXPECTED-TURNOVER": "Amount vs declared expectation",
    "TYP-CRYPTO-KEYWORD": "Crypto/virtual asset wording",
    "TYP-INDIV-PAYROLL": "Payroll-like activity on individual account",
    "TYP-CHANNEL-HOP": "Channel hopping across products",
    "TYP-TRADE-PRICING": "Trade pricing anomaly wording",
    "TYP-SENSITIVE-GOODS": "Sensitive goods reference",
    "TYP-TRAFFICKING-KEYWORD": "Trafficking/high-risk human-security wording",
    "TYP-PEP": "PEP exposure indicator",
    "TYP-SANCTIONS-JURISDICTION": "High-risk jurisdiction reference",
    "ANOM-IFOREST-CORE": "Isolation Forest anomaly scoring",
}


def _humanize_rule_id(rule_id: str) -> str:
    r = (rule_id or "").strip()
    if not r:
        return "Suspicious activity indicator"
    if r.startswith("SIM-"):
        return f"Scenario: {r[4:].replace('-', ' ').title()}"
    return _RULE_ID_TO_TITLE.get(r, r.replace("TYP-", "").replace("ANOM-", "").replace("-", " ").title())


def _red_flag_lines(alert: Dict[str, Any], enrichment: Optional[Dict[str, Any]]) -> list[Dict[str, str]]:
    why = (enrichment or {}).get("why_suspicious") if isinstance((enrichment or {}).get("why_suspicious"), dict) else {}
    typs = why.get("typologies")
    if isinstance(typs, list) and typs:
        out: list[Dict[str, str]] = []
        for x in typs:
            if not isinstance(x, dict):
                continue
            rid = str(x.get("rule_id") or "")
            title = str(x.get("title") or "").strip() or _humanize_rule_id(rid)
            nar = str(x.get("narrative") or "").strip()
            out.append({"rule_id": rid, "title": title, "narrative": nar})
        if out:
            return out
    rule_ids = alert.get("rule_ids") or []
    if isinstance(rule_ids, list) and rule_ids:
        return [
            {"rule_id": str(rid), "title": _humanize_rule_id(str(rid)), "narrative": ""}
            for rid in rule_ids
            if str(rid).strip()
        ]
    summ = str(alert.get("summary") or "").strip()
    if summ:
        return [{"rule_id": "", "title": summ[:200], "narrative": ""}]
    return [{"rule_id": "", "title": "Suspicious transaction pattern", "narrative": ""}]


def _squish_ws(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _truncate_at_boundary(text: str, max_len: int) -> str:
    t = _squish_ws(text)
    if len(t) <= max_len:
        return t
    if max_len < 4:
        return t[:max_len]
    cut = t[: max_len - 1].rsplit(" ", 1)[0]
    base = cut if cut else t[: max_len - 1]
    return base.rstrip(",;.") + "…"


def _shorten_rule_title(title: str, max_len: int = _STR_TITLE_SOFT_MAX) -> str:
    t = _squish_ws(title)
    if len(t) <= max_len:
        return t
    return _truncate_at_boundary(t, max_len)


def _red_flag_titles_compact(lines: list[Dict[str, str]]) -> str:
    raw = list(dict.fromkeys(x["title"] for x in lines if x.get("title")))
    if not raw:
        return "Suspicious transaction pattern"
    lowered = {t.lower() for t in raw}
    filtered: list[str] = []
    for t in raw:
        tl = t.lower()
        if "year-to-date inflows vs declared annual" in tl and any("unusual magnitude vs customer history" in x for x in lowered):
            continue
        if "pep exposure" in tl and any("pep - higher risk relationship" in x or "pep — higher risk relationship" in x for x in lowered):
            continue
        filtered.append(t)
    if filtered:
        raw = filtered
    titles = [_shorten_rule_title(t) for t in raw[:_STR_MAX_NATURE_TITLES]]
    base = "; ".join(titles)
    extra = len(raw) - len(titles)
    if extra > 0:
        base += f"; (+{extra} further typolog{'y' if extra == 1 else 'ies'} on file)"
    return base


def _flow_sentence_24h(
    *,
    is_outflow: bool,
    inbound_rows: list,
    outbound_rows: list,
    max_counterparties: int = _STR_TXN_FLOW_COUNTERPARTIES_MAX,
) -> str:
    mc = max(1, max_counterparties)
    parts: list[str] = []
    if not is_outflow and inbound_rows:
        bits = []
        for r in inbound_rows[:mc]:
            if not isinstance(r, dict):
                continue
            nm = str(r.get("counterparty_name") or r.get("counterparty_id") or "a third party").strip()
            amt = float(r.get("total_amount") or 0.0)
            bits.append(f"{nm} {_format_money_with_currency(amt)}")
        rest = len([x for x in inbound_rows if isinstance(x, dict)]) - len(bits)
        if bits:
            tail = f" (+{rest} more sources)" if rest > 0 else ""
            parts.append(f"Largest 24h inbound counterparty: {bits[0]}.{tail}")
    if is_outflow and outbound_rows:
        bits = []
        for r in outbound_rows[:mc]:
            if not isinstance(r, dict):
                continue
            nm = str(r.get("counterparty_name") or r.get("counterparty_id") or "a beneficiary").strip()
            amt = float(r.get("total_amount") or 0.0)
            bits.append(f"{nm} {_format_money_with_currency(amt)}")
        rest = len([x for x in outbound_rows if isinstance(x, dict)]) - len(bits)
        if bits:
            tail = f" (+{rest} more destinations)" if rest > 0 else ""
            parts.append(f"Largest 24h outbound leg: {bits[0]}.{tail}")
    return " ".join(parts)


def _safe_str(d: Dict[str, Any], *keys: str) -> str:
    for k in keys:
        v = d.get(k)
        if v is None:
            continue
        s = str(v).strip()
        if s:
            return s
    return ""


def _parse_ts_to_date(ts: Any) -> Optional[date]:
    if isinstance(ts, datetime):
        return ts.date()
    if isinstance(ts, str) and ts.strip():
        try:
            return datetime.fromisoformat(ts.replace("Z", "+00:00")).date()
        except Exception:
            pass
    return None


def _merge_paragraph_runs(paragraph) -> str:
    full = "".join(run.text for run in paragraph.runs)
    if not paragraph.runs:
        if full:
            paragraph.add_run(full)
        return full
    paragraph.runs[0].text = full
    for run in list(paragraph.runs)[1:]:
        run._element.getparent().remove(run._element)
    return full


def _is_classic_str_stationery_template(doc: Document) -> bool:
    """True when the file matches the bank's static STR layout (title + labelled customer block)."""
    paras = doc.paragraphs
    if len(paras) < 12:
        return False
    if paras[0].text.replace("\xa0", " ").strip().upper() != "SUSPICIOUS TRANSACTION REPORT":
        return False
    p1 = paras[1].text.replace("\xa0", " ").strip()
    return p1.startswith("Customer Name:")


def _replace_paragraph_text_flat(paragraph, new_text: str) -> None:
    """Set a paragraph's visible text while keeping the first run's character style."""
    _merge_paragraph_runs(paragraph)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        paragraph.runs[0].bold = False
    elif new_text:
        run = paragraph.add_run(new_text)
        run.bold = False


def _other_bvn_accounts_paragraph(customer: CustomerKyc, enrichment: Optional[Dict[str, Any]]) -> str:
    rows = (
        (enrichment or {}).get("bvn_linked_accounts")
        if isinstance((enrichment or {}).get("bvn_linked_accounts"), list)
        else []
    )
    lines: list[str] = []
    for r in rows:
        if not isinstance(r, dict):
            continue
        acct = str(r.get("account_number") or "").strip()
        if acct == customer.account_number:
            continue
        nm = str(r.get("customer_name") or "").strip() or "Account holder"
        lines.append(f"{nm} — account {acct}")
    if not lines:
        return "The bank's records confirm this BVN is linked only to this account."
    return _squish_ws(f"The bank found other accounts linked to this BVN: {'; '.join(lines)}.")


def _four_str_concern_lines(alert: Dict[str, Any], enrichment: Optional[Dict[str, Any]]) -> list[str]:
    rfl = _red_flag_lines(alert, enrichment)
    defaults = [
        "Sudden spike in activity - very high value moved through the account in a short time.",
        "Large amounts came in and went out quickly, which can indicate pass-through or layering.",
        "Money came from many different people or companies in a short period.",
        "The transfer pattern does not match the expected purpose of this account.",
    ]
    out: list[str] = []
    for i, x in enumerate(rfl[:4]):
        title = _shorten_rule_title(str(x.get("title") or "Risk indicator"))
        nar = re.sub(r"\[[A-Z0-9\-]+\]\s*", "", str(x.get("narrative") or "").strip())
        snip = _truncate_at_boundary(nar, _STR_CONCERN_BODY_MAX) if nar else ""
        body = f"{title} - {snip}" if snip else title
        out.append(f"{i + 1}. {_squish_ws(body)}")
    for d in defaults:
        if len(out) >= 4:
            break
        out.append(f"{len(out) + 1}. {d}")
    return out[:4]


def _fill_classic_str_stationery_template(doc: Document, payload: Dict[str, Any]) -> None:
    """Fill ``New STR SUSPICIOUS TRANSACTION REPORT.docx``-style stationery; paragraph-level styles are retained."""
    customer: CustomerKyc = payload["customer"]
    text: Dict[str, str] = payload["text"]
    txn: Dict[str, Any] = payload["txn"]
    alert: Dict[str, Any] = payload["alert"]
    enrichment: Optional[Dict[str, Any]] = payload.get("enrichment")
    approval = datetime.utcnow().strftime("%B %d, %Y")
    approver = payload["approver_name"].strip() or "Chief Compliance Officer"

    san = (enrichment or {}).get("sanctions_screening")
    san = san if isinstance(san, dict) else {}
    mc = int(san.get("match_count") or 0)
    ref_lists = san.get("reference_lists")
    ref_lists = ref_lists if isinstance(ref_lists, dict) else {}
    ref_san = ref_lists.get("sanctions_hits")
    ref_san = ref_san if isinstance(ref_san, list) else []
    ref_pep = ref_lists.get("pep_hits")
    ref_pep = ref_pep if isinstance(ref_pep, list) else []
    ref_am = (enrichment or {}).get("reference_adverse_media_hits")
    ref_am_l = ref_am if isinstance(ref_am, list) else []

    list_review = _squish_ws(
        f"We checked international watchlists and found no direct match to the customer's name. "
        f"However, one sanctions reference and {len(ref_am_l)} adverse-media hit(s) were flagged for analyst review. "
        f"Screening tally: watchlist {mc}, sanctions {len(ref_san)}, PEP {len(ref_pep)}, adverse {len(ref_am_l)}."
    )

    nfiu_extra = (enrichment or {}).get("nfiu_portal_note")
    if isinstance(nfiu_extra, str) and nfiu_extra.strip():
        nfiu_line = nfiu_extra.strip()[:500]
    else:
        nfiu_line = "No additional NFIU portal hit on file at compilation."

    str_supporting = _squish_ws(
        f"Over the past 12 months ({text['period_text']}), this account received {text['inflows_total_words']} "
        f"({text['inflows_total_text']}) and sent out {text['outflows_total_words']} ({text['outflows_total_text']}). "
        f"This is a very large level of movement for one account. The bank flagged this activity because money moved "
        f"suddenly and quickly, large values came in and went out within short periods, and multiple counterparties "
        f"were involved. Full details are retained in the case file."
    )
    str_supporting = _truncate_at_boundary(str_supporting, _STR_PARA_SOFT_MAX)

    flagged_amt = _format_money_with_currency(float(txn.get("amount") or 0.0))
    action_taken = _squish_ws(
        f"Action taken: The compliance team reviewed the suspicious transaction, checked 24-hour account activity, "
        f"reviewed the customer file, and completed watchlist checks. We carried out enhanced due diligence, placed "
        f"the account on closer monitoring, and updated internal records. The suspicious activity ({flagged_amt}) was "
        f"escalated through the bank's AML process, and this STR is being filed with the NFIU."
    )

    relationship_para = _squish_ws(
        f"The customer opened this account on {_date_to_long(customer.account_opened)}. "
        f"The account number is {customer.account_number}, and the BVN is {customer.id_number}."
    )

    bvn_link_para = _squish_ws(
        f"Bank records confirm this BVN belongs to this customer and is linked to account {customer.account_number}."
    )

    bvn_others = _other_bvn_accounts_paragraph(customer, enrichment)

    cdd_intro = _squish_ws(
        f"The bank's CDD/KYC record says the customer is a {customer.line_of_business}, but the transaction pattern "
        f"does not match that profile. Key concerns are listed below:"
    )

    concerns = _four_str_concern_lines(alert, enrichment)

    fills = {
        1: f"Customer Name: {customer.customer_name}",
        2: f"Customer Account Number: {customer.account_number}",
        3: f"Account Opened: {_date_to_long(customer.account_opened)}",
        4: f"Customer Address: {customer.customer_address}",
        5: f"Line of Business: {customer.line_of_business}",
        6: f"Phone Number: {customer.phone_number}",
        7: f"Date of Birth: {_date_to_long(customer.date_of_birth)}",
        8: f"ID Number: {customer.id_number}",
        9: f"Nature of Transaction: {_truncate_at_boundary(text['nature'], _STR_FIELD_NATURE_MAX)}",
        10: f"Transaction Description: {_truncate_at_boundary(text['transaction_description'], _STR_FIELD_TXN_DESC_MAX)}",
        11: _truncate_at_boundary(text["nature_clarification"], _STR_FIELD_CLARIFY_MAX),
        12: relationship_para,
        13: bvn_link_para,
        14: bvn_others,
        15: _truncate_at_boundary(list_review, _STR_PARA_SOFT_MAX),
        17: str_supporting,
        18: nfiu_line,
        19: _squish_ws(
            f"At the time of this report, no additional NFIU portal hit was identified. "
            f"Over the period {text['period_text']}, total inflow was {text['inflows_total_words']} "
            f"({text['inflows_total_text']}) and total outflow was {text['outflows_total_words']} "
            f"({text['outflows_total_text']})."
        ),
        21: cdd_intro,
        22: "\u200b",
        23: concerns[0],
        24: concerns[1],
        25: concerns[2],
        26: concerns[3],
        27: action_taken,
        31: "APPROVER:    _______________________________",
        32: approver,
        33: f"DATE:                          {approval}",
    }

    for idx, content in fills.items():
        if idx < len(doc.paragraphs):
            _replace_paragraph_text_flat(doc.paragraphs[idx], content)

    for ix in (16, 20, 28, 34, 35):
        if ix >= len(doc.paragraphs):
            continue
        raw = doc.paragraphs[ix].text.replace("\xa0", " ").replace("\u200b", "").strip()
        if not raw:
            _replace_paragraph_text_flat(doc.paragraphs[ix], "\u200b")


def _stable_rand(customer_id: str) -> int:
    h = hashlib.sha256(customer_id.encode("utf-8")).hexdigest()
    return int(h[:12], 16)


def _pick(items: list[str], idx: int) -> str:
    return items[idx % len(items)]


def _format_money(amount: float) -> str:
    # Template examples show no decimals for whole naira; keep that style.
    amt = int(round(amount))
    return f"{amt:,}"


def _format_money_2(amount: float) -> str:
    return f"{amount:,.2f}"


def _format_money_with_currency(amount: float) -> str:
    amt = int(round(amount))
    return f"₦{amt:,}"


_LESS_THAN_20 = [
    "Zero",
    "One",
    "Two",
    "Three",
    "Four",
    "Five",
    "Six",
    "Seven",
    "Eight",
    "Nine",
    "Ten",
    "Eleven",
    "Twelve",
    "Thirteen",
    "Fourteen",
    "Fifteen",
    "Sixteen",
    "Seventeen",
    "Eighteen",
    "Nineteen",
]

_TENS = {
    20: "Twenty",
    30: "Thirty",
    40: "Forty",
    50: "Fifty",
    60: "Sixty",
    70: "Seventy",
    80: "Eighty",
    90: "Ninety",
}


def _int_to_words_0_999(n: int) -> str:
    assert 0 <= n <= 999
    if n < 20:
        return _LESS_THAN_20[n]
    if n < 100:
        tens = (n // 10) * 10
        rem = n % 10
        if rem == 0:
            return _TENS[tens]
        return f"{_TENS[tens]} {_LESS_THAN_20[rem]}"
    hundreds = n // 100
    rem = n % 100
    if rem == 0:
        return f"{_LESS_THAN_20[hundreds]} Hundred"
    return f"{_LESS_THAN_20[hundreds]} Hundred and {_int_to_words_0_999(rem)}"


def _int_to_words(n: int) -> str:
    # Supports 0..999,999,999,999 (up to billions) which is enough for STR amounts.
    if n == 0:
        return "Zero"
    if n < 0:
        return f"Minus {_int_to_words(-n)}"

    billions = n // 1_000_000_000
    millions = (n // 1_000_000) % 1_000
    thousands = (n // 1_000) % 1_000
    remainder = n % 1_000

    parts: list[str] = []
    if billions:
        parts.append(f"{_int_to_words_0_999(billions)} Billion")
    if millions:
        parts.append(f"{_int_to_words_0_999(millions)} Million")
    if thousands:
        parts.append(f"{_int_to_words_0_999(thousands)} Thousand")
    if remainder:
        parts.append(_int_to_words_0_999(remainder))

    # Add commas like the sample: "Ninety Million, Fifty Five Thousand, One Hundred and ..."
    return ", ".join(parts)


def _amount_to_words(amount: float) -> str:
    total_cents = int(round(amount * 100))
    naira = total_cents // 100
    kobo = total_cents % 100
    if kobo == 0:
        return f"{_int_to_words(naira)} Naira"
    return f"{_int_to_words(naira)} Naira, {_int_to_words(kobo)} Kobo"


def _format_nigeria_phone(seed: int) -> str:
    # Generates 11-digit numbers like 08012345678
    prefixes = ["080", "081", "090", "091"]
    prefix = _pick(prefixes, seed)
    rest = seed % 10**8
    return f"{prefix}{rest:08d}"


def _format_bvn(seed: int) -> str:
    return f"{seed % 10**11:011d}"


def _format_account_number(seed: int) -> str:
    # 10-ish digits; goAML templates often accept 10–12 chars; keep simple.
    return f"{(seed % 10**10):010d}"


def _date_to_long(d: date) -> str:
    return d.strftime("%B %d, %Y")


def _month_day_year_from_seed(seed: int, start_year: int = 1980, end_year: int = 2015) -> date:
    year_range = max(1, end_year - start_year)
    year = start_year + (seed % year_range)
    month = 1 + ((seed // 7) % 12)
    # Keep day in a safe range (Word template doesn't care about exact day validity; still avoid invalid dates)
    day = 1 + ((seed // 13) % 28)
    return date(year, month, day)


@dataclass(frozen=True)
class CustomerKyc:
    customer_name: str
    account_number: str
    account_opened: date
    customer_address: str
    line_of_business: str
    phone_number: str
    date_of_birth: date
    id_number: str  # BVN / NIN-like placeholder
    customer_segment: str = "individual"  # individual | corporate
    expected_annual_turnover: Optional[float] = None
    customer_remarks: str = ""


def build_customer_kyc(
    customer_id: str,
    *,
    inferred_lob: Optional[str] = None,
    use_placeholders: bool = True,
) -> CustomerKyc:
    if use_placeholders:
        return CustomerKyc(
            customer_name="XXXXXXXXXXXX",
            account_number="XXXXXXXXXXX",
            account_opened=date(2010, 11, 19),
            customer_address="No. 5, Somorin Street, Obantoko, Abeokuta, Ogun State",
            line_of_business=inferred_lob or "Civil Servant",
            phone_number="XXXXXXXXXXXX",
            date_of_birth=date(1977, 9, 5),
            id_number="XXXXXXXXXXXXXX",
            customer_segment="individual",
            expected_annual_turnover=3_000_000.0,
            customer_remarks="",
        )
    seed = _stable_rand(customer_id)
    first_names = [
        "Amina",
        "Zainab",
        "Fatima",
        "Bola",
        "Ngozi",
        "Chinedu",
        "Ifunanya",
        "Toyin",
        "Rukayat",
        "Ibrahim",
        "Samuel",
        "Adaeze",
    ]
    last_names = [
        "Adeyemi",
        "Okafor",
        "Mohammed",
        "Lawal",
        "Abubakar",
        "Eze",
        "Oladipo",
        "Nwosu",
        "Mustapha",
        "Bello",
        "Babatunde",
        "Ibrahim",
    ]
    # Keep names human-readable but unique per customer ID in demo/synthetic generation.
    # This prevents multiple different customers from showing the same display name.
    unique_tag = f"{(seed // 101) % 10000:04d}"
    customer_name = f"{_pick(first_names, seed)} {_pick(last_names, seed // 3)} {unique_tag}"

    account_opened = _month_day_year_from_seed(seed, start_year=2003, end_year=2016)
    dob = _month_day_year_from_seed(seed // 11, start_year=1965, end_year=1995)
    bvn = _format_bvn(seed // 17)
    account_number = _format_account_number(seed // 19)
    phone_number = _format_nigeria_phone(seed // 23)

    addresses = [
        "No. 5, Somorin Street, Obantoko, Abeokuta, Ogun State",
        "12, Unity Road, GRA Ikeja, Lagos State",
        "House 3, Market Road, Kaduna North, Kaduna State",
        "Plot 8, Riverside Estate, Port Harcourt, Rivers State",
        "No. 19, Unity Avenue, Benin City, Edo State",
    ]
    customer_address = _pick(addresses, seed // 29)

    default_lob = _pick(
        ["Civil Servant", "Student", "SME Trader", "Business Owner", "Merchant", "Logistics / Importer"],
        seed,
    )
    line_of_business = inferred_lob or infer_line_of_business_from_customer_id(customer_id) or default_lob

    segment = "corporate" if (seed // 11) % 7 == 0 else "individual"
    expected_turnover = float(1_200_000 + (seed % 180) * 45_000)

    return CustomerKyc(
        customer_name=customer_name,
        account_number=account_number,
        account_opened=account_opened,
        customer_address=customer_address,
        line_of_business=line_of_business,
        phone_number=phone_number,
        date_of_birth=dob,
        id_number=bvn,
        customer_segment=segment,
        expected_annual_turnover=expected_turnover * (2.2 if segment == "corporate" else 1.0),
        customer_remarks="",
    )


def infer_line_of_business_from_txn(txn: Dict[str, Any]) -> Optional[str]:
    """Infer LOB label from transaction metadata (simulation / profile hints)."""
    inferred_lob: Optional[str] = None
    meta = txn.get("metadata") or {}
    if isinstance(meta, dict):
        profile_label = meta.get("profile") or meta.get("pattern")
        if profile_label:
            pl = str(profile_label).lower()
            if "salary" in pl or "worker" in pl or "salaried" in pl:
                inferred_lob = "Civil Servant"
            elif "student" in pl:
                inferred_lob = "Student"
            elif "trader" in pl or "sme" in pl:
                inferred_lob = "SME Trader"
            elif "hnwi" in pl or "high_net_worth" in pl or "high net" in pl:
                inferred_lob = "Business Owner"
            elif "merchant" in pl:
                inferred_lob = "Merchant"
            elif "import" in pl or "logistics" in pl:
                inferred_lob = "Logistics / Importer"
    return inferred_lob


def infer_line_of_business_from_customer_id(customer_id: str) -> Optional[str]:
    """Infer occupation/LOB from deterministic demo customer IDs."""
    cid = (customer_id or "").strip().lower()
    if not cid:
        return None
    if "student" in cid or "unilag" in cid:
        return "Student"
    if "worker" in cid or "civil" in cid or "ippis" in cid:
        return "Civil Servant"
    if "trader" in cid or "aba" in cid or "sme" in cid:
        return "SME Trader"
    if "hnwi" in cid or "business_owner" in cid:
        return "Business Owner"
    if "merchant" in cid:
        return "Merchant"
    if "importer" in cid or "apapa" in cid or "logistics" in cid:
        return "Logistics / Importer"
    return None


def _build_str_text(
    *,
    customer: CustomerKyc,
    txn: Dict[str, Any],
    alert: Dict[str, Any],
    scenario: Optional[str],
    enrichment: Optional[Dict[str, Any]] = None,
) -> Dict[str, str]:
    txn_dt = txn.get("created_at")
    if isinstance(txn_dt, datetime):
        txn_date = txn_dt.date()
    elif isinstance(txn_dt, str):
        try:
            txn_date = datetime.fromisoformat(txn_dt.replace("Z", "+00:00")).date()
        except Exception:
            txn_date = datetime.utcnow().date()
    else:
        txn_date = datetime.utcnow().date()

    amount = float(txn.get("amount") or 0.0)
    amount_currency = _format_money_with_currency(amount)
    amount_words = _amount_to_words(amount)

    en_tx = (enrichment or {}).get("transaction") if isinstance((enrichment or {}).get("transaction"), dict) else {}
    dc = str(en_tx.get("debit_credit") or "").lower()
    tx_type = str(txn.get("transaction_type") or "").lower()
    is_outflow = "outflow" in dc or "debit" in dc or "out" in tx_type

    rw = (enrichment or {}).get("rolling_windows") or {}
    h24 = rw.get("last_24_hours") if isinstance(rw.get("last_24_hours"), dict) else {}
    w_start = _parse_ts_to_date(h24.get("window_start"))
    w_end = _parse_ts_to_date(h24.get("window_end"))
    if w_start and w_end:
        window_phrase = f"in a 24-hour period from {_date_to_long(w_start)} to {_date_to_long(w_end)}"
    else:
        window_phrase = f"in a 24-hour period around {_date_to_long(txn_date)}"

    h_in = float(h24.get("inflow_total") or 0.0)
    h_out = float(h24.get("outflow_total") or 0.0)
    h_cnt = int(h24.get("transaction_count") or 0)

    ff = (enrichment or {}).get("flagged_flows") or {}
    inbound_24 = ff.get("inbound_sources_24h") if isinstance(ff.get("inbound_sources_24h"), list) else []
    outbound_24 = ff.get("outbound_destinations_24h") if isinstance(ff.get("outbound_destinations_24h"), list) else []

    best_inbound = inbound_24[0] if inbound_24 and isinstance(inbound_24[0], dict) else {}
    cp_name = _safe_str(best_inbound, "counterparty_name", "account_name", "customer_name", "counterparty_id")
    cp_bank = _safe_str(best_inbound, "bank_name", "counterparty_bank_name", "source_bank", "sender_bank")
    cp_acct = _safe_str(best_inbound, "account_number", "counterparty_account_number", "source_account_number")
    cp_amt = float(best_inbound.get("total_amount") or 0.0) if isinstance(best_inbound, dict) else 0.0
    cp_amt_txt = _format_money_with_currency(cp_amt) if cp_amt > 0 else ""
    cp_bits = [x for x in (cp_name, cp_bank, cp_acct) if x]
    cp_detail = f"{', '.join(cp_bits)}" if cp_bits else "a flagged high-risk source"
    moved_24h = h_in + h_out

    transaction_description_24h = _squish_ws(
        f"{window_phrase}, this customer's account ({customer.account_number}, BVN {customer.id_number}) showed "
        f"unusual activity. Within this period, the account received {_amount_to_words(h_in)} "
        f"({_format_money_with_currency(h_in)}) and sent out {_amount_to_words(h_out)} "
        f"({_format_money_with_currency(h_out)}) across {h_cnt or 'several'} transactions. "
        f"The largest incoming transfer was {cp_amt_txt or 'a high-value amount'} from {cp_detail}. "
        f"The bank also marked a specific posting of {amount_words} ({amount_currency}) as suspicious."
    )

    rflines = _red_flag_lines(alert, enrichment)
    red_flag_titles = _red_flag_titles_compact(rflines)

    h72 = rw.get("last_72_hours") if isinstance(rw.get("last_72_hours"), dict) else {}
    h72_in = float(h72.get("inflow_total") or 0.0)
    h72_out = float(h72.get("outflow_total") or 0.0)
    nature_clarification = _squish_ws(
        f"To add more detail: In one day, over {_format_money_with_currency(moved_24h)} moved in and out of this "
        f"customer's account. This is a very large amount changing hands very quickly, which is unusual. "
        f"Over the past three days, the account recorded credits of {_format_money_with_currency(h72_in)} and debits "
        f"of {_format_money_with_currency(h72_out)}. This pattern may indicate rapid pass-through or layering."
    )

    narrative = (txn.get("narrative") or "").strip()

    inflows_total = float(alert.get("inflows_total") or 0.0)
    outflows_total = float(alert.get("outflows_total") or 0.0)
    inflows_text = f"₦{_format_money_2(inflows_total)}"
    outflows_text = f"₦{_format_money_2(outflows_total)}"
    inflows_words = _amount_to_words(inflows_total)
    outflows_words = _amount_to_words(outflows_total)
    period_text = alert.get("period_text") or "January 1, 2025, to March 13, 2026"

    rule_ids = alert.get("rule_ids") or []
    rule_joined = " ".join(str(r) for r in rule_ids).upper() if isinstance(rule_ids, list) else str(rule_ids).upper()
    suspicion_summary = (
        "Large outflows, rapid fund movement, and inconsistent transaction patterns."
        if is_outflow
        else "Large inflows, wire spikes, and inconsistent transaction patterns."
    )
    if "STRUCTUR" in rule_joined:
        suspicion_summary = "Repeated deposits, structuring-like spacing, and inconsistent transaction patterns."
    elif "SMURF" in rule_joined:
        suspicion_summary = "Multiple inbound transfers, smurfing indicators, and inconsistent transaction patterns."
    elif scenario and "LAYER" in str(scenario).upper():
        suspicion_summary = "Large inflows, suspected layering / pass-through activity, and rapid onward transfers."

    title_list = list(
        dict.fromkeys(
            _shorten_rule_title(str(x.get("title") or "").strip())
            for x in rflines
            if str(x.get("title") or "").strip()
        )
    )
    shown = title_list[:_STR_EVIDENCE_TITLE_ONLY_MAX]
    evidence_body = "; ".join(shown) if shown else suspicion_summary
    if len(title_list) > len(shown):
        evidence_body += f" (+{len(title_list) - len(shown)} more on file)"

    return {
        "nature": red_flag_titles,
        "transaction_description": transaction_description_24h,
        "nature_clarification": nature_clarification,
        "red_flag_explanation": nature_clarification,
        "inflows_total_text": f"{inflows_text}",
        "outflows_total_text": f"{outflows_text}",
        "inflows_total_words": inflows_words,
        "outflows_total_words": outflows_words,
        "txn_date_long": _date_to_long(txn_date),
        "txn_amount_words": amount_words,
        "period_text": period_text,
        "narrative": narrative,
        "is_outflow": is_outflow,
        "suspicion_summary": suspicion_summary,
        "str_evidence_rules": _truncate_at_boundary(evidence_body, 360),
    }


def _assemble_str_doc_payload(
    *,
    customer: CustomerKyc,
    txn: Dict[str, Any],
    alert: Dict[str, Any],
    approver_name: str,
    enrichment: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    scenario = None
    rule_ids = alert.get("rule_ids") or []
    if isinstance(rule_ids, list):
        for rid in rule_ids:
            if isinstance(rid, str) and rid.startswith("SIM-"):
                scenario = rid.replace("SIM-", "")
                break

    text = _build_str_text(
        customer=customer,
        txn=txn,
        alert=alert,
        scenario=scenario,
        enrichment=enrichment,
    )

    return {
        "customer": customer,
        "txn": txn,
        "alert": alert,
        "enrichment": enrichment,
        "text": text,
        "approver_name": approver_name,
        "scenario": scenario,
    }


_STR_WORD_TEMPLATE_ERR = (
    'STR Word template not found. Place "New STR SUSPICIOUS TRANSACTION REPORT.docx" at the repository root '
    'or under backend/demo_assets, or set AML_STR_WORD_TEMPLATE to the file path.'
)


def render_str_docx_bytes(
    *,
    customer: CustomerKyc,
    txn: Dict[str, Any],
    alert: Dict[str, Any],
    approver_name: str,
    enrichment: Optional[Dict[str, Any]] = None,
) -> bytes:
    """
    Populate the bank ``New STR SUSPICIOUS TRANSACTION REPORT.docx`` stationery only (paragraphs are filled in
    place so fonts and bolding from that file are preserved). No alternate Roman-numeral layout is generated.
    """
    template_path = resolve_str_word_template_path()
    if not template_path:
        raise FileNotFoundError(_STR_WORD_TEMPLATE_ERR)

    payload = _assemble_str_doc_payload(
        customer=customer,
        txn=txn,
        alert=alert,
        approver_name=approver_name,
        enrichment=enrichment,
    )

    doc = Document(str(template_path))
    if not _is_classic_str_stationery_template(doc):
        raise ValueError(
            "The STR Word file is not the expected New STR stationery layout "
            f"(title 'SUSPICIOUS TRANSACTION REPORT' with a Customer Name block): {template_path}"
        )

    _fill_classic_str_stationery_template(doc, payload)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()

